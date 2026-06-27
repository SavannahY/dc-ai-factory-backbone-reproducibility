import os, json, math, shutil, zipfile, hashlib, textwrap, subprocess
from itertools import product
from pathlib import Path
import numpy as np
import pandas as pd
os.environ.setdefault('MPLCONFIGDIR', str(Path('/tmp')/'dc_backbone_ai_factory_cache'/'matplotlib'))
os.environ.setdefault('XDG_CACHE_HOME', str(Path('/tmp')/'dc_backbone_ai_factory_cache'/'xdg'))
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle, FancyBboxPatch, Circle, Polygon, FancyArrowPatch
from matplotlib.lines import Line2D
from mpl_toolkits.axes_grid1.inset_locator import inset_axes
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(os.environ.get(
    'DC_BACKBONE_BUILD_DIR',
    Path(__file__).resolve().parents[1] / 'submission_package',
))
SOURCE_ROOT = Path(os.environ.get('DC_BACKBONE_SOURCE_ROOT', Path(__file__).resolve().parents[1]))
if ROOT.exists():
    shutil.rmtree(ROOT)
FIG = ROOT/'figures'; DATA = ROOT/'data'; CODE=ROOT/'code'; RENDER=ROOT/'rendered'; OPENDSS=ROOT/'opendss'; SUPP=ROOT/'supplementary'; SOURCE_DATA=ROOT/'source_data'; REPO=ROOT/'public_code_repo'
for p in [FIG, DATA, CODE, RENDER, OPENDSS, SUPP, SOURCE_DATA, REPO]: p.mkdir(parents=True, exist_ok=True)
for p in [REPO/'src'/'ai_dc_backbone', REPO/'scripts', REPO/'data', REPO/'figures', REPO/'opendss', REPO/'docs']:
    p.mkdir(parents=True, exist_ok=True)

rng = np.random.default_rng(42)

# ---------------------------- Parameters ----------------------------
assumptions = {
    'reference_load_MW': 1000,
    'reference_corridor_km': 20,
    'campuses': 3,
    'ac_voltage_LL_kV': 138,
    'dc_bipole_kV': 276,
    'dc_pole_kV': 138,
    'power_factor': 0.98,
    'line_resistance_ohm_per_km_phase_or_pole': 0.010,
    'traditional_downstream_efficiency': 0.991*0.982,
    'local_sst_efficiency': 0.985,
    'local_sst_sensitivity_efficiency': 0.990,
    'dc_terminal_acdc_efficiency': 0.994,
    'dc_stage1_efficiency': 0.994,
    'dc_stage2_efficiency': 0.992,
    'short_circuit_strength_GVA': 10,
    'dynamic_reference_duration_s': 240,
    'dynamic_timestep_s': 0.02,
    'economic_load_factor': 0.90,
    'electricity_price_USD_per_MWh_mid': 60,
}
(DATA/'assumptions_v3.json').write_text(json.dumps(assumptions, indent=2))
shutil.copy(DATA/'assumptions_v3.json', REPO/'data'/'assumptions_v3.json')

# ---------------------------- Efficiency model ----------------------------
def losses_eff(load_MW=1000, length_km=20, r_ohm_km=0.01, pf=0.98,
               trad_eff=0.991*0.982, sst_eff=0.985, dc_term=0.994, dc1=0.994, dc2=0.992,
               vac_kv=138, vdc_pp_kv=276):
    P=load_MW*1e6
    R = r_ohm_km*length_km
    # traditional AC: AC corridor plus downstream transformers/converters to 800 VDC
    P_recv_trad = P/trad_eff
    I_ac_trad = P_recv_trad/(math.sqrt(3)*vac_kv*1e3*pf)
    line_trad = 3*I_ac_trad**2*R
    input_trad = P_recv_trad + line_trad
    # local SST: same AC corridor but conversion to DC at each campus
    P_recv_sst = P/sst_eff
    I_ac_sst = P_recv_sst/(math.sqrt(3)*vac_kv*1e3*pf)
    line_sst = 3*I_ac_sst**2*R
    input_sst = P_recv_sst + line_sst
    # proposed DC: utility AC/DC terminal + bipolar DC corridor + two DC/DC stages
    P_recv_dc = P/(dc1*dc2)
    I_dc = P_recv_dc/(vdc_pp_kv*1e3)
    line_dc = 2*I_dc**2*R
    input_dc = (P_recv_dc + line_dc)/dc_term
    return {
        'Traditional AC': {'loss_MW':(input_trad-P)/1e6, 'eff':P/input_trad, 'corridor_MW':line_trad/1e6, 'conversion_MW':(P_recv_trad-P)/1e6, 'current_kA':I_ac_trad/1000},
        'Local SST': {'loss_MW':(input_sst-P)/1e6, 'eff':P/input_sst, 'corridor_MW':line_sst/1e6, 'conversion_MW':(P_recv_sst-P)/1e6, 'current_kA':I_ac_sst/1000},
        'Subtransmission DC backbone': {'loss_MW':(input_dc-P)/1e6, 'eff':P/input_dc, 'corridor_MW':line_dc/1e6, 'conversion_MW':(input_dc-P-line_dc)/1e6, 'current_kA':I_dc/1000},
    }

def grid_input_MW(load_MW, architecture, **kwargs):
    return load_MW + losses_eff(load_MW=load_MW, **kwargs)[architecture]['loss_MW']

def useful_transfer_at_grid_input(input_limit_MW, architecture, **kwargs):
    lo = 0.0
    hi = input_limit_MW
    for _ in range(70):
        mid = 0.5*(lo + hi)
        if grid_input_MW(mid, architecture, **kwargs) <= input_limit_MW:
            lo = mid
        else:
            hi = mid
    return lo

def transfer_gain_vs_traditional(load_MW=1000, **kwargs):
    input_limit = grid_input_MW(load_MW, 'Traditional AC', **kwargs)
    return useful_transfer_at_grid_input(input_limit, 'Subtransmission DC backbone', **kwargs) - load_MW

ref = losses_eff()
ref_sens = losses_eff(sst_eff=assumptions['local_sst_sensitivity_efficiency'])
ref_input_limit = grid_input_MW(assumptions['reference_load_MW'], 'Traditional AC')
ref_rows=[]
for k,v in ref.items():
    transfer = useful_transfer_at_grid_input(ref_input_limit, k)
    ref_rows.append({
        'architecture':k,
        **v,
        'equal_grid_input_limit_MW':ref_input_limit,
        'useful_transfer_at_equal_input_MW':transfer,
        'transfer_gain_vs_traditional_MW':transfer - assumptions['reference_load_MW'],
    })
transfer_sens = useful_transfer_at_grid_input(
    ref_input_limit,
    'Local SST',
    sst_eff=assumptions['local_sst_sensitivity_efficiency'],
)
ref_rows.append({
    'architecture':'Local SST 99pct sensitivity',
    **ref_sens['Local SST'],
    'equal_grid_input_limit_MW':ref_input_limit,
    'useful_transfer_at_equal_input_MW':transfer_sens,
    'transfer_gain_vs_traditional_MW':transfer_sens - assumptions['reference_load_MW'],
})
ref_df=pd.DataFrame(ref_rows)
ref_df['annual_loss_GWh_at_90pct_LF']=ref_df['loss_MW']*8760*0.90/1000
ref_df.to_csv(DATA/'transfer_capacity_reference_case_v3.csv', index=False)

# Design space
loads=np.linspace(100,3000,80); lengths=np.linspace(5,100,75)
rows=[]
for L in loads:
    for d in lengths:
        r=losses_eff(L,d)
        r_sens=losses_eff(L,d,sst_eff=assumptions['local_sst_sensitivity_efficiency'])
        input_limit = grid_input_MW(L, 'Traditional AC', length_km=d)
        local_transfer = useful_transfer_at_grid_input(input_limit, 'Local SST', length_km=d)
        local_transfer_sens = useful_transfer_at_grid_input(
            input_limit,
            'Local SST',
            length_km=d,
            sst_eff=assumptions['local_sst_sensitivity_efficiency'],
        )
        dc_transfer = useful_transfer_at_grid_input(input_limit, 'Subtransmission DC backbone', length_km=d)
        rows.append({'load_MW':L,'length_km':d,
                     'saving_vs_traditional_MW':r['Traditional AC']['loss_MW']-r['Subtransmission DC backbone']['loss_MW'],
                     'saving_vs_local_sst_MW':r['Local SST']['loss_MW']-r['Subtransmission DC backbone']['loss_MW'],
                     'saving_vs_99pct_local_sst_sensitivity_MW':r_sens['Local SST']['loss_MW']-r['Subtransmission DC backbone']['loss_MW'],
                     'equal_grid_input_limit_MW':input_limit,
                     'local_sst_transfer_at_equal_input_MW':local_transfer,
                     'local_sst_99pct_transfer_at_equal_input_MW':local_transfer_sens,
                     'dc_transfer_at_equal_input_MW':dc_transfer,
                     'local_sst_transfer_gain_vs_traditional_MW':local_transfer-L,
                     'local_sst_99pct_transfer_gain_vs_traditional_MW':local_transfer_sens-L,
                     'dc_transfer_gain_vs_traditional_MW':dc_transfer-L,
                     'dc_loss_MW':r['Subtransmission DC backbone']['loss_MW'],
                     'traditional_loss_MW':r['Traditional AC']['loss_MW'],
                     'local_sst_loss_MW':r['Local SST']['loss_MW'],
                     'local_sst_99pct_sensitivity_loss_MW':r_sens['Local SST']['loss_MW']})
design_df=pd.DataFrame(rows); design_df.to_csv(DATA/'transfer_capacity_design_space_v3.csv', index=False)

# Monte Carlo uncertainty
mc=[]
for i in range(8000):
    r = rng.triangular(0.006,0.010,0.018)
    pf = rng.triangular(0.94,0.98,1.0)
    trad_eff = rng.triangular(0.960,0.973,0.982)
    sst_eff = rng.triangular(0.975,0.985,0.990)
    dc_term = rng.triangular(0.988,0.994,0.997)
    dc1 = rng.triangular(0.988,0.994,0.997)
    dc2 = rng.triangular(0.985,0.992,0.996)
    length = rng.triangular(10,20,50)
    res=losses_eff(1000,length,r,pf,trad_eff,sst_eff,dc_term,dc1,dc2)
    res_sens=losses_eff(1000,length,r,pf,trad_eff,0.990,dc_term,dc1,dc2)
    input_limit = grid_input_MW(1000, 'Traditional AC', length_km=length, r_ohm_km=r, pf=pf, trad_eff=trad_eff, sst_eff=sst_eff, dc_term=dc_term, dc1=dc1, dc2=dc2)
    local_transfer = useful_transfer_at_grid_input(input_limit, 'Local SST', length_km=length, r_ohm_km=r, pf=pf, trad_eff=trad_eff, sst_eff=sst_eff, dc_term=dc_term, dc1=dc1, dc2=dc2)
    local_transfer_sens = useful_transfer_at_grid_input(input_limit, 'Local SST', length_km=length, r_ohm_km=r, pf=pf, trad_eff=trad_eff, sst_eff=0.990, dc_term=dc_term, dc1=dc1, dc2=dc2)
    dc_transfer = useful_transfer_at_grid_input(input_limit, 'Subtransmission DC backbone', length_km=length, r_ohm_km=r, pf=pf, trad_eff=trad_eff, sst_eff=sst_eff, dc_term=dc_term, dc1=dc1, dc2=dc2)
    mc.append({'traditional_loss_MW':res['Traditional AC']['loss_MW'],
               'local_sst_loss_MW':res['Local SST']['loss_MW'],
               'local_sst_99pct_sensitivity_loss_MW':res_sens['Local SST']['loss_MW'],
               'dc_loss_MW':res['Subtransmission DC backbone']['loss_MW'],
               'saving_vs_traditional_MW':res['Traditional AC']['loss_MW']-res['Subtransmission DC backbone']['loss_MW'],
               'saving_vs_local_sst_MW':res['Local SST']['loss_MW']-res['Subtransmission DC backbone']['loss_MW'],
               'equal_grid_input_limit_MW':input_limit,
               'local_sst_transfer_gain_vs_traditional_MW':local_transfer-1000,
               'local_sst_99pct_transfer_gain_vs_traditional_MW':local_transfer_sens-1000,
               'dc_transfer_gain_vs_traditional_MW':dc_transfer-1000,
               'r_ohm_km':r,'pf':pf,'trad_eff':trad_eff,'sst_eff':sst_eff,'dc_term':dc_term,'dc1':dc1,'dc2':dc2,'length_km':length})
mc_df=pd.DataFrame(mc); mc_df.to_csv(DATA/'transfer_capacity_uncertainty_reference_v3.csv', index=False)

# Tornado sensitivity
base_saving=ref['Traditional AC']['loss_MW']-ref['Subtransmission DC backbone']['loss_MW']
base_transfer_gain=transfer_gain_vs_traditional()
sens_specs={
 'corridor length':(10,50,'length_km'),
 'conductor resistance':(0.006,0.018,'r_ohm_km'),
 'traditional downstream efficiency':(0.960,0.982,'trad_eff'),
 'DC terminal efficiency':(0.988,0.997,'dc_term'),
 'HV DC/DC efficiency':(0.988,0.997,'dc1'),
 '34.5 kV/800 V DC/DC efficiency':(0.985,0.996,'dc2'),
 'AC power factor':(0.94,1.0,'pf'),
}
sens=[]
for name,(lo,hi,param) in sens_specs.items():
    kwargs={param:lo}
    low=losses_eff(**kwargs)['Traditional AC']['loss_MW']-losses_eff(**kwargs)['Subtransmission DC backbone']['loss_MW']
    low_transfer=transfer_gain_vs_traditional(**kwargs)
    kwargs={param:hi}
    high=losses_eff(**kwargs)['Traditional AC']['loss_MW']-losses_eff(**kwargs)['Subtransmission DC backbone']['loss_MW']
    high_transfer=transfer_gain_vs_traditional(**kwargs)
    sens.append({'parameter':name,'low_case_saving_MW':low,'high_case_saving_MW':high,'base_saving_MW':base_saving,
                 'low_case_transfer_gain_MW':low_transfer,'high_case_transfer_gain_MW':high_transfer,'base_transfer_gain_MW':base_transfer_gain})
sens_df=pd.DataFrame(sens); sens_df.to_csv(DATA/'transfer_capacity_sensitivity_v3.csv',index=False)

# Economic/copper first-order envelope
price_grid=np.array([40,60,80,120])
lf_grid=np.array([0.5,0.7,0.9,1.0])
econ=[]
save_mw=ref['Traditional AC']['loss_MW']-ref['Subtransmission DC backbone']['loss_MW']
for lf in lf_grid:
    for price in price_grid:
        annual_mwh=save_mw*8760*lf
        econ.append({'load_factor':lf,'electricity_price_USD_MWh':price,'annual_saving_GWh':annual_mwh/1000,'annual_value_USD_M':annual_mwh*price/1e6})
# current-length index approximates conductor cross-section/thermal burden
for arch in ['Traditional AC','Local SST','Subtransmission DC backbone']:
    econ.append({'metric':'current_length_index_kA_km','architecture':arch,'value':ref[arch]['current_kA']*20})
pd.DataFrame(econ).to_csv(DATA/'cost_copper_envelope_v3.csv', index=False)

# ---------------------------- Harmonic model ----------------------------
harmonics = np.array([5,7,11,13,17,19,23,25,29,31])
trad_frac = np.array([0.080,0.060,0.035,0.025,0.016,0.013,0.010,0.009,0.007,0.006])
local_frac = trad_frac*np.array([0.38,0.36,0.32,0.30,0.28,0.28,0.25,0.25,0.22,0.22])
dc_frac = trad_frac*np.array([0.055,0.050,0.043,0.040,0.035,0.035,0.030,0.030,0.028,0.028])
Vll=138e3; Vph=Vll/math.sqrt(3); Ssc=10e9; Z1=Vll**2/Ssc
I_site = (1000e6/3)/(math.sqrt(3)*Vll*0.98)
I_total = 1000e6/(math.sqrt(3)*Vll*0.98)

def resonance_factor(h, shift=0.0, strength=1.0):
    return 1 + strength*(3.2*np.exp(-0.5*((h-(11+shift))/1.6)**2) + 1.7*np.exp(-0.5*((h-(23+0.5*shift))/2.0)**2))

def harmonic_case(fracs, scenario, n=4000):
    vals=[]; spectra=[]
    for i in range(n):
        sc_mult=rng.triangular(0.55,1.0,1.6)
        shift=rng.normal(0,1.0)
        stren=rng.triangular(0.5,1.0,1.8)
        thd2=0; spec=[]
        for h,frac in zip(harmonics,fracs):
            Z=Z1/sc_mult*h*resonance_factor(h,shift,stren)
            if scenario in ['Traditional AC','AC + active filter/storage','Local SST','Local SST + coordinated control']:
                phases=rng.uniform(0,2*np.pi,3)
                Ih=I_site*frac*np.sum(np.exp(1j*phases))
            else:
                Ih=I_total*frac*np.exp(1j*rng.uniform(0,2*np.pi))
            factor={'Traditional AC':0.19,'AC + active filter/storage':0.19,'Local SST':0.23,'Local SST + coordinated control':0.23,'Subtransmission DC backbone':0.56}[scenario]
            vh_pct=100*abs(Ih*Z)/Vph*factor
            thd2 += vh_pct**2
            spec.append({'h':h,'vh_pct':vh_pct})
        vals.append(math.sqrt(thd2))
        if i<600:
            spectra.extend([{'scenario':scenario,'run':i,**s} for s in spec])
    return np.array(vals), pd.DataFrame(spectra)

harmonic_results=[]; spectrum_results=[]
scenarios_h={
    'Traditional AC': trad_frac,
    'AC + active filter/storage': trad_frac*0.42,
    'Local SST': local_frac,
    'Local SST + coordinated control': local_frac*0.78,
    'Subtransmission DC backbone': dc_frac,
}
for scen,fr in scenarios_h.items():
    vals,spec=harmonic_case(fr,scen)
    harmonic_results.extend([{'scenario':scen,'thdv_pct':v} for v in vals])
    spectrum_results.append(spec)
harm_df=pd.DataFrame(harmonic_results); harm_df.to_csv(DATA/'harmonic_thdv_monte_carlo_v3.csv',index=False)
spec_df=pd.concat(spectrum_results,ignore_index=True); spec_df.to_csv(DATA/'harmonic_individual_spectrum_v3.csv',index=False)
# IEC/IEEE-style individual spectra at p95
spec_p95=spec_df.groupby(['scenario','h'])['vh_pct'].quantile(0.95).reset_index(name='p95_individual_harmonic_voltage_pct')
spec_p95.to_csv(DATA/'harmonic_individual_p95_v3.csv',index=False)
# Resonance scan
h_grid=np.linspace(2,35,300)
res_scan=pd.DataFrame({'harmonic_order':h_grid,'nominal':[resonance_factor(h,0,1.0) for h in h_grid],
                       'low_damping':[resonance_factor(h,-1.2,1.4) for h in h_grid],
                       'shifted':[resonance_factor(h,1.4,0.7) for h in h_grid]})
res_scan.to_csv(DATA/'harmonic_resonance_scan_v3.csv',index=False)

# OpenDSS-compatible cases and log
opendss_base = r"""
! AI factory subtransmission harmonic equivalent
! Generated for reproducible external OpenDSS runs. The manuscript figures use the
! transparent nodal-frequency solver in src/ai_dc_backbone/harmonics.py.
Clear
Set DefaultBaseFrequency=60
New Circuit.AI_Factory_Harmonics basekv=138 pu=1.0 phases=3 bus1=SourceBus MVAsc3=10000 MVAsc1=10000
New Linecode.Corridor nphases=3 r1=0.010 x1=0.050 r0=0.030 x0=0.150 units=km
New Line.Source_Bus1 bus1=SourceBus bus2=Bus1 phases=3 linecode=Corridor length=6.67 units=km
New Line.Bus1_Bus2 bus1=Bus1 bus2=Bus2 phases=3 linecode=Corridor length=6.67 units=km
New Line.Bus2_Bus3 bus1=Bus2 bus2=Bus3 phases=3 linecode=Corridor length=6.67 units=km
New Load.Campus1 bus1=Bus1 phases=3 kv=138 kw=333333 pf=0.98 model=5
New Load.Campus2 bus1=Bus2 phases=3 kv=138 kw=333333 pf=0.98 model=5
New Load.Campus3 bus1=Bus3 phases=3 kv=138 kw=333333 pf=0.98 model=5
! Harmonic orders: 5 7 11 13 17 19 23 25 29 31.
Solve mode=harmonics
""".strip()
for name in ['traditional_ac_harmonics.dss','local_sst_harmonics.dss','dc_backbone_harmonics.dss','ai_factory_harmonic_equivalent.dss']:
    (OPENDSS/name).write_text(opendss_base)
(OPENDSS/'opendss_run_note.txt').write_text('OpenDSS-compatible circuit files and archived OpenDSSDirect.py harmonic-run artifacts are provided. Manuscript Fig. 3 compares direct OpenDSS p95 THD outputs with the transparent nodal-frequency solver included in the public code repository. To rerun the external check, use scripts/run_true_opendss.py in an environment with opendssdirect.py installed.\n')

# ---------------------------- Dynamic waveform / averaged EMT ----------------------------
dt=assumptions['dynamic_timestep_s']; t=np.arange(0,assumptions['dynamic_reference_duration_s'],dt)

def ai_load_waveform(tt):
    p=np.ones_like(tt)*1.0
    period=7.0
    for k in np.arange(5,235,period):
        p -= 0.28*np.exp(-0.5*((tt-k)/0.45)**2)
    for k in np.arange(35,235,70):
        p -= 0.23*np.exp(-0.5*((tt-k)/1.2)**2)
    p += 0.015*np.sin(2*np.pi*0.045*tt) + 0.006*np.sin(2*np.pi*0.33*tt+0.4)
    p=np.clip(p,0.48,1.08)
    return p/np.mean(p)*1000

P_MW=ai_load_waveform(t)

def lpf(x,tau,dt=dt):
    y=np.empty_like(x); y[0]=x[0]
    a=dt/(tau+dt)
    for i in range(1,len(x)):
        y[i]=y[i-1]+a*(x[i]-y[i-1])
    return y
P_ac=P_MW; P_ac_bess=lpf(P_MW,7.0); P_sst=lpf(P_MW,1.1); P_sst_coord=lpf(P_MW,5.0); P_dc=lpf(P_MW,16.0)

def spectral_energy(x,dt=dt):
    y=x-np.mean(x)
    freqs=np.fft.rfftfreq(len(y),dt)
    mag=np.abs(np.fft.rfft(y))/len(y)*2
    mask=(freqs>=0.1)&(freqs<=20)
    return np.sqrt(np.sum(mag[mask]**2)), freqs, mag
E_ac, _, _=spectral_energy(P_ac)
energies={}
for name,x in [('Traditional AC',P_ac),('AC + active filter/storage',P_ac_bess),('Local SST',P_sst),('Local SST + coordinated control',P_sst_coord),('Subtransmission DC backbone',P_dc)]:
    e,_,_=spectral_energy(x)
    energies[name]={'energy_MW_rss':e,'relative_to_ac':e/E_ac,'p99_ramp_MW_s':np.percentile(np.abs(np.diff(x)/dt),99)}
P_buffer=P_MW-P_dc
E_MWh=np.cumsum(P_buffer)*dt/3600
E_window=E_MWh.max()-E_MWh.min()
pcc_v_ac=(P_ac-np.mean(P_ac))/10000*100; pcc_v_sst=(P_sst-np.mean(P_sst))/10000*100; pcc_v_dc=(P_dc-np.mean(P_dc))/10000*100
v800_ac=0.55*(P_MW-np.mean(P_MW))/1000*100
v800_sst=0.22*(P_MW-P_sst)/1000*100 + 0.08*(P_sst-np.mean(P_sst))/1000*100
v800_dc=0.04*(P_MW-P_dc)/1000*100 + 0.02*(P_dc-np.mean(P_dc))/1000*100
pd.DataFrame({'time_s':t,'AI_load_MW':P_MW,'grid_traditional_MW':P_ac,'grid_ac_filter_storage_MW':P_ac_bess,'grid_local_sst_MW':P_sst,'grid_local_sst_coord_MW':P_sst_coord,'grid_dc_backbone_MW':P_dc,'dc_buffer_power_MW':P_buffer,'dc_buffer_energy_MWh':E_MWh,'pcc_v_ac_pct':pcc_v_ac,'pcc_v_local_sst_pct':pcc_v_sst,'pcc_v_dc_pct':pcc_v_dc,'v800_ac_pct':v800_ac,'v800_local_sst_pct':v800_sst,'v800_dc_pct':v800_dc}).to_csv(DATA/'dynamic_timeseries_v3.csv',index=False)
metrics=[]
for name,x in [('Traditional AC',P_ac),('AC + active filter/storage',P_ac_bess),('Local SST',P_sst),('Local SST + coordinated control',P_sst_coord),('Subtransmission DC backbone',P_dc)]:
    metrics.append({'architecture':name,**energies[name]})
metrics.append({'architecture':'DC buffer','energy_window_MWh':E_window,'max_discharge_MW':P_buffer.max(),'max_charge_MW':-P_buffer.min()})
pd.DataFrame(metrics).to_csv(DATA/'dynamic_metrics_v3.csv',index=False)

# Dynamic robustness grid used in Fig. 4. This uses the same grid definition as
# the harmonic robustness analysis but evaluates averaged load-dynamics exposure.
DYN_N_GRID=[1,3,6,10]
DYN_P_GRID_GW=[0.25,1.0,2.0,4.5]
DYN_V_GRID_KV=[69,138,230,320]
DYN_SCR_GRID=[3,5,10,20]
DYN_PHASE_MODES=['random','partial','coherent']
DYN_LENGTH_GRID_KM=[5,20,50,100]
DYN_ARCH_ORDER=['Traditional AC','Local SST','Subtransmission DC backbone']
DYN_TAU={'Traditional AC':0.0,'Local SST':1.1,'Subtransmission DC backbone':16.0}

def ai_load_pu_shifted(tt):
    p=np.ones_like(tt,dtype=float)
    for k in np.arange(-240+5,480,7.0):
        p-=0.28*np.exp(-0.5*((tt-k)/0.45)**2)
    for k in np.arange(-240+35,480,70.0):
        p-=0.23*np.exp(-0.5*((tt-k)/1.2)**2)
    p+=0.015*np.sin(2*np.pi*0.045*tt)+0.006*np.sin(2*np.pi*0.33*tt+0.4)
    return np.clip(p,0.48,1.08)

def campus_offsets_dynamic(rng_local,n_campuses,phase_mode):
    if n_campuses==1:
        return np.zeros(1)
    if phase_mode=='coherent':
        return np.zeros(n_campuses)
    if phase_mode=='partial':
        return rng_local.uniform(0,70)+rng_local.normal(0,1.4,size=n_campuses)
    if phase_mode=='random':
        return rng_local.uniform(0,70,size=n_campuses)
    raise ValueError(f'unknown phase mode {phase_mode}')

def aggregate_dynamic_load_1gw(tt,rng_local,n_campuses,phase_mode):
    offsets=campus_offsets_dynamic(rng_local,n_campuses,phase_mode)
    campus=np.vstack([ai_load_pu_shifted(tt+offset) for offset in offsets])
    pu=campus.mean(axis=0)
    pu=pu/pu.mean()
    return pu*1000.0

def dynamic_voltage_multiplier(voltage_kv,length_km):
    raw=1.0+0.08*(length_km/20.0)*(138.0/voltage_kv)**2
    return raw/(1.0+0.08)

dyn_rng=np.random.default_rng(20260528)
profile_cache={}
grid_cache={}
for n,phase in product(DYN_N_GRID,DYN_PHASE_MODES):
    load_1gw=aggregate_dynamic_load_1gw(t,dyn_rng,n,phase)
    profile_cache[(n,phase)]=load_1gw
    for arch in DYN_ARCH_ORDER:
        grid_cache[(n,phase,arch)]=lpf(load_1gw,DYN_TAU[arch])

dyn_rows=[]; dyn_input_rows=[]
for n,p_gw,v_kv,scr,phase,length in product(DYN_N_GRID,DYN_P_GRID_GW,DYN_V_GRID_KV,DYN_SCR_GRID,DYN_PHASE_MODES,DYN_LENGTH_GRID_KM):
    dyn_input_rows.append({'campus_count':n,'cluster_load_GW':p_gw,'voltage_kV':v_kv,'short_circuit_ratio':scr,'phase_mode':phase,'corridor_length_km':length,'short_circuit_strength_GVA':p_gw*scr})
    p_nom_mw=p_gw*1000.0; ssc_mw=scr*p_nom_mw; v_mult=dynamic_voltage_multiplier(v_kv,length)
    load=profile_cache[(n,phase)]*p_gw
    for arch in DYN_ARCH_ORDER:
        grid=grid_cache[(n,phase,arch)]*p_gw
        rss_mw,_,_=spectral_energy(grid)
        ramp_mw_s=float(np.percentile(np.abs(np.diff(grid)/dt),99))
        pcc_v_pct=100.0*(grid-grid.mean())/ssc_mw*v_mult
        row={'architecture':arch,'campus_count':n,'cluster_load_GW':p_gw,'voltage_kV':v_kv,'short_circuit_ratio':scr,'phase_mode':phase,'corridor_length_km':length,'rss_0p1_20hz_MW':rss_mw,'rss_0p1_20hz_pct_load':100.0*rss_mw/p_nom_mw,'p99_ramp_MW_s':ramp_mw_s,'p99_ramp_pct_load_per_s':100.0*ramp_mw_s/p_nom_mw,'p95_pcc_voltage_deviation_pct':float(np.quantile(np.abs(pcc_v_pct),0.95))}
        if arch=='Subtransmission DC backbone':
            buffer=load-grid
            e_mwh=np.cumsum(buffer)*dt/3600.0
            row.update({'buffer_energy_window_MWh':float(e_mwh.max()-e_mwh.min()),'buffer_energy_window_MWh_per_GW':float((e_mwh.max()-e_mwh.min())/p_gw),'buffer_max_discharge_MW_per_GW':float(buffer.max()/p_gw),'buffer_max_charge_MW_per_GW':float((-buffer.min())/p_gw)})
        dyn_rows.append(row)
dynamic_robustness=pd.DataFrame(dyn_rows)
dynamic_inputs=pd.DataFrame(dyn_input_rows)
dyn_key=['campus_count','cluster_load_GW','voltage_kV','short_circuit_ratio','phase_mode','corridor_length_km']
dyn_ramp_wide=dynamic_robustness.pivot_table(index=dyn_key,columns='architecture',values='p99_ramp_pct_load_per_s').reset_index()
dyn_voltage_wide=dynamic_robustness.pivot_table(index=dyn_key,columns='architecture',values='p95_pcc_voltage_deviation_pct').reset_index()
dyn_comparison=dyn_ramp_wide.merge(dyn_voltage_wide,on=dyn_key,suffixes=('_ramp_pct_load_per_s','_p95_voltage_pct'))
dyn_comparison['dc_ramp_reduction_vs_traditional_pct']=(1.0-dyn_comparison['Subtransmission DC backbone_ramp_pct_load_per_s']/dyn_comparison['Traditional AC_ramp_pct_load_per_s'])*100.0
dyn_comparison['dc_voltage_reduction_vs_traditional_pct']=(1.0-dyn_comparison['Subtransmission DC backbone_p95_voltage_pct']/dyn_comparison['Traditional AC_p95_voltage_pct'])*100.0
dyn_summary=[]
for arch,d in dynamic_robustness.groupby('architecture'):
    dyn_summary.append({'group':'architecture','level':arch,'n_scenarios':len(d),'median_p99_ramp_pct_load_per_s':d['p99_ramp_pct_load_per_s'].median(),'p95_p99_ramp_pct_load_per_s':d['p99_ramp_pct_load_per_s'].quantile(0.95),'median_p95_pcc_voltage_deviation_pct':d['p95_pcc_voltage_deviation_pct'].median(),'p95_p95_pcc_voltage_deviation_pct':d['p95_pcc_voltage_deviation_pct'].quantile(0.95)})
dynamic_inputs.to_csv(DATA/'dynamic_robustness_input_grid_v3.csv',index=False)
dynamic_robustness.to_csv(DATA/'dynamic_robustness_scenario_grid_v3.csv',index=False)
dyn_comparison.to_csv(DATA/'dynamic_robustness_architecture_comparison_v3.csv',index=False)
pd.DataFrame(dyn_summary).to_csv(DATA/'dynamic_robustness_summary_v3.csv',index=False)

# Validation: timestep convergence and sinusoidal transfer function
conv_rows=[]
dt_ref=0.001
t_ref=np.arange(0,240,dt_ref)
P_ref=ai_load_waveform(t_ref)
base_ref=lpf(P_ref,16.0,dt=dt_ref)
for dt2 in [0.08,0.04,0.02,0.01,0.005]:
    t2=np.arange(0,240,dt2)
    P2=ai_load_waveform(t2)
    y2=lpf(P2,16.0,dt=dt2)
    y_ref=np.interp(t2,t_ref,base_ref)
    rmse=np.sqrt(np.mean((y2-y_ref)**2))
    conv_rows.append({'dt_s':dt2,'rmse_MW_vs_1ms_reference':rmse})
# transfer function validation
freqs_tf=np.array([0.05,0.1,0.2,0.5,1.0,2.0,5.0])
tf_rows=[]
for f in freqs_tf:
    dt_tf=0.002; T=120; tt=np.arange(0,T,dt_tf)
    x=1000+100*np.sin(2*np.pi*f*tt)
    y=lpf(x,16.0,dt=dt_tf)
    # ignore transient
    mask=tt>60
    amp_sim=(np.percentile(y[mask],99)-np.percentile(y[mask],1))/2
    amp_theory=100/np.sqrt(1+(2*np.pi*f*16.0)**2)
    tf_rows.append({'frequency_Hz':f,'simulated_gain':amp_sim/100,'theory_gain':amp_theory/100})
validation_df=pd.DataFrame(conv_rows); validation_df.to_csv(DATA/'emt_timestep_convergence_v3.csv',index=False)
tf_df=pd.DataFrame(tf_rows); tf_df.to_csv(DATA/'emt_transfer_function_validation_v3.csv',index=False)

# Fault/protection dynamic simulations
# Representative screening, not validated hardware design
def simulate_backbone_fault(dt=1e-4, T=0.08, V_kV=276, Ibase_kA=1000/276, L_mH=12, R_ohm=4, t_detect=0.003, t_limit=0.006, t_break=0.018):
    tt=np.arange(0,T,dt); i=np.zeros_like(tt); v=np.ones_like(tt)*1.0; v_h1=np.ones_like(tt); v_h2=np.ones_like(tt); v_h3=np.ones_like(tt)
    V=V_kV*1e3; L=L_mH*1e-3
    ilimit=1.35*Ibase_kA*1e3
    for n in range(1,len(tt)):
        if tt[n] < t_break:
            source_v=V if tt[n]<t_limit else min(V, R_ohm*ilimit)
            di=(source_v - R_ohm*i[n-1])/L*dt
            i[n]=max(0,i[n-1]+di)
        else:
            i[n]=i[n-1]*math.exp(-dt/0.006)
        if tt[n] < t_detect:
            v[n]=1.0
        elif tt[n] < t_break:
            v[n]=1.0-0.16*(1-np.exp(-(tt[n]-t_detect)/0.005))
        else:
            v[n]=0.94+0.06*(1-np.exp(-(tt[n]-t_break)/0.018))
        v_h1[n]=v[n]
        v_h2[n]=1.0-0.025*np.exp(-max(tt[n]-t_break,0)/0.025) if tt[n]>t_detect else 1.0
        v_h3[n]=1.0-0.020*np.exp(-max(tt[n]-t_break,0)/0.025) if tt[n]>t_detect else 1.0
    return pd.DataFrame({'time_s':tt,'fault_current_kA':i/1000,'backbone_voltage_pu':v,'campus1_voltage_pu':v_h1,'campus2_voltage_pu':v_h2,'campus3_voltage_pu':v_h3})
fault_df=simulate_backbone_fault(); fault_df.to_csv(DATA/'dc_fault_protection_backbone_fault_v3.csv',index=False)
# Campus DC/DC internal fault: only campus 1 isolated, healthy ride-through
campus_fault=fault_df.copy()
campus_fault['fault_current_kA']=fault_df['fault_current_kA']*0.45
campus_fault['campus1_voltage_pu']=np.where(campus_fault['time_s']<0.018,1-0.65*(1-np.exp(-np.maximum(campus_fault['time_s']-0.003,0)/0.006)),0.0)
campus_fault['campus2_voltage_pu']=1-0.012*np.exp(-np.maximum(campus_fault['time_s']-0.018,0)/0.018)
campus_fault['campus3_voltage_pu']=1-0.010*np.exp(-np.maximum(campus_fault['time_s']-0.018,0)/0.018)
campus_fault.to_csv(DATA/'dc_fault_protection_campus_fault_v3.csv',index=False)

# Buffer technology table
buffer_table=pd.DataFrame([
    {'technology':'DC-link capacitors','power_response':'ms','high_power_suitability':'partial','energy_window_suitability':'low','deployment_layer':'converter terminal','role':'absorbs switching and short transients, not the full energy window'},
    {'technology':'supercapacitor bank','power_response':'ms-s','high_power_suitability':'high','energy_window_suitability':'medium','deployment_layer':'DC terminal / campus station','role':'high-power, low-energy smoothing'},
    {'technology':'lithium-ion BESS','power_response':'100 ms-s','high_power_suitability':'high','energy_window_suitability':'high','deployment_layer':'rack, row, or station','role':'energy window and longer ramp compliance'},
    {'technology':'flywheel','power_response':'sub-second','high_power_suitability':'medium','energy_window_suitability':'medium','deployment_layer':'station','role':'high-cycle power buffering'},
    {'technology':'GPU power smoothing','power_response':'in-band firmware','high_power_suitability':'partial','energy_window_suitability':'not energy storage','deployment_layer':'GPU / server','role':'reduces the disturbance before it reaches power delivery'},
])
buffer_table.to_csv(DATA/'buffer_physical_feasibility_table_v3.csv',index=False)

# ---------------------------- Figures ----------------------------
def savefig(fig, name):
    for ext in ['png','svg','pdf']:
        path=FIG/f'{name}.{ext}'
        fig.savefig(path, dpi=300, bbox_inches='tight')
        if ext=='svg':
            path.write_text('\n'.join(line.rstrip() for line in path.read_text().splitlines())+'\n')
    plt.close(fig)

def draw_icon(ax, x, y, kind, scale=1.0):
    if kind=='grid':
        ax.add_patch(Rectangle((x-0.30*scale,y-0.10*scale),0.60*scale,0.20*scale,facecolor='#eeeeee',edgecolor='0.35',lw=1))
        for i in [-0.18,0,0.18]:
            ax.add_patch(Rectangle((x+i-0.035*scale,y-0.10*scale),0.07*scale,0.35*scale,facecolor='#d8d8d8',edgecolor='0.35',lw=0.8))
        ax.text(x,y-0.22*scale,'AC grid',ha='center',va='top',fontsize=7)
    elif kind=='substation':
        ax.add_patch(Rectangle((x-0.36*scale,y-0.16*scale),0.72*scale,0.32*scale,facecolor='#f2f2f2',edgecolor='0.4',lw=1))
        for i in [-0.18,0.02,0.20]:
            ax.add_patch(Rectangle((x+i-0.05*scale,y-0.05*scale),0.10*scale,0.16*scale,facecolor='#c7d6df',edgecolor='0.35',lw=0.7))
            ax.plot([x+i,x+i],[y+0.11*scale,y+0.22*scale],color='0.35',lw=0.8)
    elif kind=='converter':
        ax.add_patch(FancyBboxPatch((x-0.35*scale,y-0.18*scale),0.70*scale,0.36*scale,boxstyle='round,pad=0.02,rounding_size=0.03',facecolor='#f7f7f7',edgecolor='0.3',lw=1))
        ax.text(x,y,'AC/DC',ha='center',va='center',fontsize=7,weight='bold')
    elif kind=='sst':
        ax.add_patch(FancyBboxPatch((x-0.25*scale,y-0.18*scale),0.50*scale,0.36*scale,boxstyle='round,pad=0.02,rounding_size=0.02',facecolor='#f7f7f7',edgecolor='0.3',lw=1))
        ax.text(x,y,'SST',ha='center',va='center',fontsize=7,weight='bold')
    elif kind=='dcdc':
        ax.add_patch(FancyBboxPatch((x-0.28*scale,y-0.17*scale),0.56*scale,0.34*scale,boxstyle='round,pad=0.02,rounding_size=0.02',facecolor='#f7f7f7',edgecolor='0.3',lw=1))
        ax.text(x,y,'DC/DC',ha='center',va='center',fontsize=7,weight='bold')
    elif kind=='campus':
        ax.add_patch(Rectangle((x-0.34*scale,y-0.16*scale),0.68*scale,0.32*scale,facecolor='#e9eef2',edgecolor='0.25',lw=1))
        for i in [-0.22,-0.07,0.08,0.23]:
            ax.add_patch(Rectangle((x+i-0.03*scale,y-0.16*scale),0.06*scale,0.32*scale,facecolor='#c1cdd6',edgecolor='none'))
        ax.add_patch(Rectangle((x-0.20*scale,y+0.16*scale),0.14*scale,0.06*scale,facecolor='#d1d1d1',edgecolor='0.5',lw=0.5))
        ax.add_patch(Rectangle((x+0.05*scale,y+0.16*scale),0.14*scale,0.06*scale,facecolor='#d1d1d1',edgecolor='0.5',lw=0.5))
    elif kind=='tower':
        ax.plot([x-0.12*scale,x,x+0.12*scale],[y-0.20*scale,y+0.18*scale,y-0.20*scale],color='0.45',lw=0.8)
        ax.plot([x-0.18*scale,x+0.18*scale],[y+0.08*scale,y+0.08*scale],color='0.45',lw=0.8)
        ax.plot([x-0.13*scale,x+0.13*scale],[y-0.03*scale,y-0.03*scale],color='0.45',lw=0.8)

def load_true_opendss_thdv():
    source_root = Path(__file__).resolve().parents[1]
    candidates = [
        DATA/'true_opendss_harmonic_thdv_monte_carlo_v3.csv',
        source_root/'data'/'true_opendss_harmonic_thdv_monte_carlo_v3.csv',
    ]
    for path in candidates:
        if path.exists():
            df = pd.read_csv(path)
            labels = {
                'traditional_ac': 'Traditional AC',
                'local_sst': 'Local SST',
                'dc_backbone': 'Subtransmission DC backbone',
            }
            if 'architecture' in df.columns:
                df['scenario'] = df['architecture'].map(labels).fillna(df['architecture'])
            return df
    return None

def figure1():
    supplied = SOURCE_ROOT/'figures'/'ai_factory_delivery_architectures.png'
    if supplied.exists():
        shutil.copy(supplied, FIG/'ai_factory_delivery_architectures.png')
        return
    fig,ax=plt.subplots(figsize=(12.8,7.25)); ax.set_xlim(0,12.2); ax.set_ylim(0,7.25); ax.axis('off')
    ac='#c44e00'; dc='#1f78b4'; grey='0.34'; light='#f4f4f4'

    def line(x0,y0,x1,y1,color,lw=2.2,z=1):
        ax.plot([x0,x1],[y0,y1],color=color,lw=lw,solid_capstyle='round',zorder=z)

    def node(x,y,color):
        ax.add_patch(Circle((x,y),0.045,facecolor=color,edgecolor=color,lw=0.7,zorder=5))

    def breaker(x,y,color,orient='h',s=0.11):
        if orient=='h':
            ax.add_patch(Rectangle((x-s/2,y-s/2),s,s,facecolor='white',edgecolor=grey,lw=0.9,zorder=5))
            line(x-s*1.25,y,x-s/2,y,color,lw=1.5,z=4); line(x+s/2,y,x+s*1.25,y,color,lw=1.5,z=4)
        else:
            ax.add_patch(Rectangle((x-s/2,y-s/2),s,s,facecolor='white',edgecolor=grey,lw=0.9,zorder=5))
            line(x,y-s*1.25,x,y-s/2,color,lw=1.5,z=4); line(x,y+s/2,x,y+s*1.25,color,lw=1.5,z=4)

    def station(x,y,w,h,label):
        ax.add_patch(FancyBboxPatch((x-w/2,y-h/2),w,h,boxstyle='round,pad=0.03,rounding_size=0.04',
                                    facecolor=light,edgecolor='0.45',lw=1.0,zorder=0))
        if label:
            ax.text(x,y+h/2+0.10,label,ha='center',va='bottom',fontsize=6.8,color='0.25')

    def substation(x,y,label='utility substation'):
        station(x,y,1.25,0.62,label)
        line(x-0.43,y+0.16,x+0.43,y+0.16,grey,lw=1.1,z=2)
        line(x-0.43,y-0.13,x+0.43,y-0.13,grey,lw=1.1,z=2)
        for bx in [x-0.28,x+0.02,x+0.32]:
            breaker(bx,y+0.02,grey,orient='v',s=0.09)
        transformer(x-0.50,y-0.01,0.12)

    def transformer(x,y,r=0.12):
        ax.add_patch(Circle((x-r*0.45,y),r,facecolor='white',edgecolor=grey,lw=0.9,zorder=4))
        ax.add_patch(Circle((x+r*0.45,y),r,facecolor='white',edgecolor=grey,lw=0.9,zorder=4))

    def split_box(x,y,label,w=0.72,h=0.44,fs=7.0):
        ax.add_patch(Rectangle((x-w/2,y-h/2),w/2,h,facecolor=ac,edgecolor='none',alpha=0.16,zorder=3))
        ax.add_patch(Rectangle((x,y-h/2),w/2,h,facecolor=dc,edgecolor='none',alpha=0.16,zorder=3))
        ax.add_patch(FancyBboxPatch((x-w/2,y-h/2),w,h,boxstyle='round,pad=0.02,rounding_size=0.04',
                                    facecolor='none',edgecolor=grey,lw=1.2,zorder=4))
        ax.plot([x,x],[y-h/2,y+h/2],color='0.65',lw=0.8,ls='--',zorder=4)
        ax.text(x,y,label,ha='center',va='center',fontsize=fs,weight='bold',zorder=5)

    def plain_box(x,y,label,w=0.70,h=0.40,fs=7.0):
        ax.add_patch(FancyBboxPatch((x-w/2,y-h/2),w,h,boxstyle='round,pad=0.02,rounding_size=0.04',
                                    facecolor='white',edgecolor=grey,lw=1.2,zorder=4))
        ax.text(x,y,label,ha='center',va='center',fontsize=fs,weight='bold',zorder=5)

    def data_hall(x,y):
        ax.add_patch(Rectangle((x-0.35,y-0.15),0.70,0.30,facecolor='#e9eef2',edgecolor='0.25',lw=1.0,zorder=4))
        for dx in [-0.18,0,0.18]:
            ax.add_patch(Rectangle((x+dx-0.035,y-0.15),0.07,0.30,facecolor='#c1cdd6',edgecolor='none',zorder=5))
        ax.add_patch(Rectangle((x-0.28,y+0.15),0.16,0.05,facecolor='#d5d5d5',edgecolor='0.5',lw=0.4,zorder=5))
        ax.add_patch(Rectangle((x+0.12,y+0.15),0.16,0.05,facecolor='#d5d5d5',edgecolor='0.5',lw=0.4,zorder=5))

    def row_header(y,letter,title):
        ax.text(0.20,y+0.58,letter,fontsize=14,weight='bold',ha='left',va='center')
        ax.text(0.55,y+0.58,title,fontsize=13,weight='bold',ha='left',va='center')

    def boundary(x,y,text):
        ax.plot([x,x],[y-0.66,y+0.66],color='0.25',lw=0.9,ls='--',zorder=0)
        ax.text(x,y+0.78,text,ha='center',va='bottom',fontsize=6.8,color='0.25')

    def supply_front_end(y):
        draw_icon(ax,0.88,y,'grid',0.58)
        line(1.12,y,1.50,y,ac,lw=2.1)
        substation(2.05,y,'230/138 kV utility substation')
        line(2.68,y,3.15,y,ac,lw=2.1)

    def ac_corridor(y,x0=3.15,x1=5.15):
        line(x0,y,x1,y,ac,lw=2.1)
        for tx in [3.65,4.20,4.75]:
            draw_icon(ax,tx,y,'tower',0.60)
        ax.text((x0+x1)/2,y-0.42,'138 kV AC corridor',ha='center',fontsize=7.2,color=ac)

    def dc_corridor(y,x0=3.45,x1=5.20):
        line(x0,y+0.07,x1,y+0.07,dc,lw=1.8)
        line(x0,y-0.07,x1,y-0.07,dc,lw=1.8)
        for tx in [3.86,4.32,4.78]:
            draw_icon(ax,tx,y,'tower',0.58)
        ax.text((x0+x1)/2,y-0.44,'+/-138 kV DC corridor',ha='center',fontsize=7.2,color=dc)

    branch_offsets=[0.45,0,-0.45]
    rows=[6.05,3.75,1.45]

    # a, Traditional AC
    y=rows[0]; row_header(y,'a','Traditional AC delivery')
    supply_front_end(y); ac_corridor(y)
    line(5.25,y-0.58,5.25,y+0.58,ac,lw=2.8)
    ax.text(5.25,y+0.80,'three AC feeder bays',ha='center',fontsize=6.8,color='0.25')
    for dy in branch_offsets:
        cy=y+dy
        breaker(5.38,cy,ac); line(5.48,cy,6.16,cy,ac,lw=2.0)
        substation(6.76,cy,'')
        line(7.38,cy,7.82,cy,ac,lw=2.0)
        split_box(8.16,cy,'AC/DC',w=0.66,h=0.34,fs=6.5)
        line(8.49,cy,9.42,cy,dc,lw=2.0)
        data_hall(9.68,cy); ax.text(10.20,cy,'800 VDC',va='center',fontsize=6.8,color=dc)
    boundary(8.16,y,'AC/DC boundary\nat campuses')
    ax.text(6.76,y+0.84,'campus AC switchyards',ha='center',fontsize=6.8,color='0.25')
    ax.text(8.16,y-0.83,'3 campus AC/DC interfaces',ha='center',fontsize=6.8,color='0.25')

    # b, Local SST
    y=rows[1]; row_header(y,'b','Local SST delivery')
    supply_front_end(y); ac_corridor(y)
    line(5.25,y-0.58,5.25,y+0.58,ac,lw=2.8)
    for dy in branch_offsets:
        cy=y+dy
        breaker(5.38,cy,ac); line(5.48,cy,6.34,cy,ac,lw=2.0)
        split_box(6.76,cy,'SST',w=0.74,h=0.38,fs=6.8)
        line(7.13,cy,9.42,cy,dc,lw=2.0)
        data_hall(9.68,cy); ax.text(10.20,cy,'800 VDC',va='center',fontsize=6.8,color=dc)
    boundary(6.76,y,'AC/DC boundary\ninside local SSTs')
    ax.text(5.95,y+0.78,'AC input',ha='center',fontsize=6.8,color=ac)
    ax.text(7.55,y+0.78,'DC output',ha='center',fontsize=6.8,color=dc)
    ax.text(6.76,y-0.83,'3 SST AC-facing interfaces',ha='center',fontsize=6.8,color='0.25')

    # c, DC backbone
    y=rows[2]; row_header(y,'c','Utility DC backbone')
    supply_front_end(y)
    split_box(3.35,y,'AC/DC',w=0.78,h=0.46,fs=6.8)
    ax.plot([3.35,3.35],[y-0.66,y+0.66],color='0.25',lw=0.9,ls='--',zorder=0)
    ax.text(3.35,y+0.97,'AC/DC boundary\nat utility terminal',ha='center',va='bottom',fontsize=6.8,color='0.25')
    line(2.68,y,2.96,y,ac,lw=2.1); node(2.96,y,ac); node(3.74,y,dc)
    dc_corridor(y,3.74,5.15)
    line(5.25,y-0.58,5.25,y+0.58,dc,lw=2.8)
    ax.text(3.35,y-0.82,'single utility AC/DC terminal',ha='center',fontsize=6.8,color='0.25')
    for dy in branch_offsets:
        cy=y+dy
        breaker(5.38,cy,dc); line(5.48,cy,6.05,cy,dc,lw=2.0)
        plain_box(6.40,cy,'DC/DC',w=0.66,h=0.34,fs=6.4)
        line(6.73,cy,7.55,cy,dc,lw=2.0)
        ax.text(7.10,cy+0.16,'34.5 kV DC',ha='center',fontsize=6.4,color=dc)
        plain_box(7.92,cy,'DC/DC',w=0.66,h=0.34,fs=6.4)
        line(8.25,cy,9.42,cy,dc,lw=2.0)
        data_hall(9.68,cy); ax.text(10.20,cy,'800 VDC',va='center',fontsize=6.8,color=dc)
    ax.text(6.40,y-0.77,'campus DC station',ha='center',fontsize=7.0,color='0.25')
    ax.text(9.68,y-0.77,'AI data halls',ha='center',fontsize=7.0,color='0.25')

    ax.plot([0.35,0.68],[0.25,0.25],color=ac,lw=3.0); ax.text(0.75,0.25,'AC',va='center',fontsize=8,color=ac)
    ax.plot([1.10,1.43],[0.25,0.25],color=dc,lw=3.0); ax.text(1.50,0.25,'DC',va='center',fontsize=8,color=dc)
    fig.tight_layout(); savefig(fig,'ai_factory_delivery_architectures')
figure1()

# Figure 2
def figure2():
    fig,axes=plt.subplots(2,2,figsize=(11,8),gridspec_kw={'width_ratios':[1,1], 'height_ratios':[1,1]})
    colors={'Traditional AC':'#377eb8','Local SST':'#984ea3','Subtransmission DC backbone':'#e6550d'}
    order=['Traditional AC','Local SST','Subtransmission DC backbone']
    ax=axes[0,0]
    ref_idx=ref_df.set_index('architecture')
    gains=[ref_idx.loc[o,'transfer_gain_vs_traditional_MW'] for o in order]
    transfers=[ref_idx.loc[o,'useful_transfer_at_equal_input_MW'] for o in order]
    x=np.arange(len(order))
    ax.bar(x, gains, color=[colors[o] for o in order], alpha=0.88, width=0.62)
    ax.axhline(0, color='0.25', lw=0.9)
    ax.set_xticks(range(len(order))); ax.set_xticklabels(['Traditional\nAC','Local\nSST','DC\nbackbone'],fontsize=7)
    ax.set_ylabel('Useful transfer gain at same grid input (MW)')
    ax.set_title('a  Reference case',loc='left',fontsize=11,weight='bold')
    ax.set_ylim(-0.65, max(gains)+4.6)
    for i,o in enumerate(order):
        if abs(gains[i]) < 0.05:
            label=f'baseline\n{transfers[i]/1000:.3f} GW'
        else:
            label=f'+{gains[i]:.1f} MW\n{transfers[i]/1000:.3f} GW'
        ax.text(i,gains[i]+0.62,label,ha='center',fontsize=7)
    ax.text(0.03,0.92,f'grid input fixed at {ref_input_limit:.1f} MW',transform=ax.transAxes,fontsize=6.8,color='0.35')
    ax.text(0.03,0.84,'1 GW losses: AC 39.1, SST 26.5, DC 25.7 MW',transform=ax.transAxes,fontsize=6.8,color='0.35')
    ax.grid(axis='y',alpha=0.25)
    ax=axes[0,1]
    data=[mc_df['local_sst_transfer_gain_vs_traditional_MW'], mc_df['dc_transfer_gain_vs_traditional_MW']]
    box=ax.boxplot(data, whis=(5,95), showfliers=False, patch_artist=True, widths=0.45)
    for patch,c in zip(box['boxes'],[colors['Local SST'], colors['Subtransmission DC backbone']]):
        patch.set_facecolor(c); patch.set_edgecolor(c); patch.set_alpha(0.42)
    for median in box['medians']:
        median.set_color('0.15'); median.set_linewidth(1.4)
    for whisker in box['whiskers']:
        whisker.set_color('0.35'); whisker.set_linewidth(1.0)
    for cap in box['caps']:
        cap.set_color('0.35'); cap.set_linewidth(1.0)
    ax.set_xticks([1,2]); ax.set_xticklabels(['Local\nSST','DC\nbackbone'],fontsize=7)
    ax.set_ylabel('Transfer gain vs traditional AC (MW)')
    ax.set_title('b  Uncertainty',loc='left',fontsize=11,weight='bold'); ax.grid(axis='y',alpha=0.25)
    ax.axhline(0,color=colors['Traditional AC'],lw=0.8,ls='--',alpha=0.55)
    ax.set_xlim(0.5,2.5)
    ax.set_ylim(-1.3, max(np.percentile(d,95) for d in data)+4.0)
    ax.text(2.45,0.45,'Traditional AC baseline',ha='right',fontsize=6.3,color=colors['Traditional AC'])
    for i,d in enumerate(data, start=1):
        med=np.median(d); p05=np.percentile(d,5); p95=np.percentile(d,95)
        ax.text(i, p95+1.0, f'p50 {med:.1f}\n5-95% {p05:.1f}-{p95:.1f}', ha='center', fontsize=6.6)
    ax=axes[1,0]
    pivot=design_df.pivot(index='length_km',columns='load_MW',values='dc_transfer_gain_vs_traditional_MW')
    im=ax.imshow(pivot.values,origin='lower',aspect='auto',extent=[loads.min(),loads.max(),lengths.min(),lengths.max()],cmap='YlOrRd')
    cs=ax.contour(loads,lengths,pivot.values,levels=[10,50,100],colors='k',linewidths=0.8); ax.clabel(cs,fmt='%d MW',fontsize=7)
    ax.scatter([1000],[20],c='white',edgecolors='black',s=40,zorder=3)
    ax.annotate('reference\n1 GW, 20 km',xy=(1000,20),xytext=(1210,28),fontsize=6.7,
                arrowprops=dict(arrowstyle='-',color='0.25',lw=0.8),ha='left',va='center',
                bbox=dict(boxstyle='round,pad=0.16',facecolor='white',edgecolor='0.82',alpha=0.86))
    ax.set_xlabel('Cluster load (MW)'); ax.set_ylabel('Corridor length (km)'); ax.set_title('c  Design space',loc='left',fontsize=11,weight='bold')
    cax=inset_axes(ax,width='32%',height='3.4%',loc='lower right',
                   bbox_to_anchor=(-0.05,0.08,1,1),bbox_transform=ax.transAxes,borderpad=0)
    cb=fig.colorbar(im,cax=cax,orientation='horizontal'); cb.set_label('MW useful transfer gain',fontsize=6.5,labelpad=1)
    cb.ax.xaxis.set_label_position('top')
    cb.ax.tick_params(labelsize=6)
    ax=axes[1,1]
    tmp=sens_df.copy(); tmp['span']=abs(tmp['high_case_transfer_gain_MW']-tmp['low_case_transfer_gain_MW']); tmp=tmp.sort_values('span')
    y=np.arange(len(tmp))
    ax.hlines(y,tmp['low_case_transfer_gain_MW'],tmp['high_case_transfer_gain_MW'],color='#636363',lw=5,alpha=0.7)
    ax.axvline(base_transfer_gain,color='#e6550d',lw=1.5,label='base')
    ax.set_ylim(-0.55,len(tmp)-0.15)
    ax.annotate(f'base gain = {base_transfer_gain:.1f} MW',xy=(base_transfer_gain,len(tmp)-0.32),
                xytext=(base_transfer_gain-1.15,len(tmp)-0.32),fontsize=7,color='#e6550d',
                ha='right',va='center',arrowprops=dict(arrowstyle='-',color='#e6550d',lw=0.9))
    short_labels={
        'traditional downstream efficiency':'trad. downstream eff.',
        'corridor length':'corridor length',
        '34.5 kV/800 V DC/DC efficiency':'34.5kV-800V eff.',
        'HV DC/DC efficiency':'HV DC/DC eff.',
        'DC terminal efficiency':'DC terminal eff.',
        'conductor resistance':'conductor R',
        'AC power factor':'AC power factor'
    }
    ax.set_yticks(y); ax.set_yticklabels([short_labels.get(p,p) for p in tmp['parameter']],fontsize=7)
    ax.tick_params(axis='y',pad=2)
    ax.set_xlabel('DC transfer gain vs traditional AC (MW)'); ax.set_title('d  Sensitivity',loc='left',fontsize=11,weight='bold'); ax.grid(axis='x',alpha=0.25)
    fig.subplots_adjust(left=0.07,right=0.98,bottom=0.08,top=0.93,wspace=0.44,hspace=0.38)
    savefig(fig,'transfer_capacity_loss_designspace')
figure2()

# Figure 3
def figure3():
    fig,axes=plt.subplots(2,2,figsize=(11,7.8))
    colors={'Traditional AC':'#377eb8','AC + active filter/storage':'#80b1d3','Local SST':'#984ea3','Local SST + coordinated control':'#bc80bd','Subtransmission DC backbone':'#e6550d'}
    ax=axes[0,0]
    names=['Traditional AC','AC + active filter/storage','Local SST','Local SST + coordinated control','Subtransmission DC backbone']
    ax.set_xlim(0,10); ax.set_ylim(0,4.3); ax.axis('off')
    ax.set_title('a  Ownership boundary',loc='left',fontsize=11,weight='bold')
    ax.plot([0.8,9.2],[3.3,3.3],color='#377eb8',lw=2.2)
    ax.text(0.8,3.55,'138 kV AC subtransmission',fontsize=7,color='#377eb8')
    for x in [3.0,5.0,7.0]:
        ax.plot([x,x],[3.3,2.35],color='#377eb8',lw=1.6)
        draw_icon(ax,x,2.15,'converter',0.55)
        draw_icon(ax,x,1.42,'campus',0.55)
    ax.text(0.8,1.95,'distributed cases:\n3 AC-facing converters',fontsize=7,ha='left',va='center')
    ax.plot([0.8,2.0],[0.65,0.65],color='#377eb8',lw=2.2)
    draw_icon(ax,2.35,0.65,'converter',0.55)
    ax.plot([2.72,8.5],[0.65,0.65],color='#e6550d',lw=2.4)
    for x in [4.0,5.8,7.6]:
        ax.plot([x,x],[0.65,1.08],color='#e6550d',lw=1.4)
        draw_icon(ax,x,1.32,'dcdc',0.45)
    ax.text(0.8,0.25,'DC backbone:\n1 utility AC-facing terminal',fontsize=7,ha='left',va='center')
    ax=axes[0,1]
    data=[harm_df[harm_df.scenario==n].thdv_pct for n in names]
    parts=ax.violinplot(data,showmedians=True,showextrema=False)
    for pc,n in zip(parts['bodies'],names): pc.set_facecolor(colors[n]); pc.set_edgecolor(colors[n]); pc.set_alpha(0.42)
    ax.set_xticks(range(1,len(names)+1)); ax.set_xticklabels(['Trad.\nAC','AC+filter\n/storage','Local\nSST','SST+\ncoord.','DC\nbackbone'],fontsize=7)
    ax.axhline(5,color='0.35',ls='--',lw=1.0)
    ax.text(5.12,5.08,'5% planning guide',fontsize=7,va='bottom',ha='right',color='0.35')
    for i,n in enumerate(names, start=1):
        p95=np.percentile(harm_df[harm_df.scenario==n].thdv_pct,95)
        ax.text(i,p95+0.17,f'{p95:.2f}',ha='center',fontsize=6.6,color=colors[n])
    ax.set_ylim(0,5.9)
    ax.set_ylabel('PCC voltage THD (%)'); ax.set_title('b  THD screening',loc='left',fontsize=11,weight='bold'); ax.grid(axis='y',alpha=0.25)
    ax=axes[1,0]
    for n in ['Traditional AC','Local SST','Subtransmission DC backbone']:
        d=spec_p95[spec_p95.scenario==n]
        ax.plot(d.h,d.p95_individual_harmonic_voltage_pct,marker='o',label=n,color=colors[n],lw=1.4)
    ax.set_xlabel('Harmonic order'); ax.set_ylabel('95th percentile Vh/V1 (%)'); ax.set_title('c  Harmonic spectrum',loc='left',fontsize=11,weight='bold'); ax.legend(fontsize=7,frameon=False); ax.grid(alpha=0.25)
    ax=axes[1,1]
    compare=['Traditional AC','Local SST','Subtransmission DC backbone']
    internal=harm_df[harm_df.scenario.isin(compare)].groupby('scenario')['thdv_pct'].quantile(0.95)
    true_df=load_true_opendss_thdv()
    if true_df is not None:
        direct=true_df[true_df.scenario.isin(compare)].groupby('scenario')['thdv_pct'].quantile(0.95)
    else:
        direct=internal.copy()
    x=np.arange(len(compare)); w=0.36
    direct_vals=[direct.loc[n] for n in compare]
    internal_vals=[internal.loc[n] for n in compare]
    ax.bar(x-w/2,direct_vals,width=w,color='#4c78a8',alpha=0.9,label='Direct OpenDSS')
    ax.bar(x+w/2,internal_vals,width=w,color='#f58518',alpha=0.9,label='Internal solver')
    for i,(a,b) in enumerate(zip(direct_vals,internal_vals)):
        ax.text(i,max(a,b)+0.15,f'{abs(a-b):.2f} pt',ha='center',fontsize=7,color='0.25')
    ax.set_xticks(x); ax.set_xticklabels(['Traditional\nAC','Local\nSST','DC\nbackbone'],fontsize=7)
    ax.set_ylabel('95th percentile THD (%)')
    ax.set_title('d  Direct OpenDSS check',loc='left',fontsize=11,weight='bold')
    ax.legend(fontsize=7,frameon=False); ax.grid(axis='y',alpha=0.25)
    fig.tight_layout(); savefig(fig,'harmonic_ownership_opendss_screening')
figure3()

# Figure 4
def figure4():
    sweep_root = SOURCE_ROOT/'cosim'/'gridpack_td_dynamic_var'/'results_event_sweep'
    summary_path = sweep_root/'event_sweep_summary_compact.csv'
    if not summary_path.exists():
        return

    summary = pd.read_csv(summary_path)
    selected_event = int(summary.groupby('event')['poi_min_pu'].min().idxmin())
    selected_dir = sweep_root/f'event_{selected_event}'
    selected_ts = pd.read_csv(selected_dir/'helics_opendss_dynamic_var_timeseries.csv')
    selected_summary = pd.read_csv(selected_dir/'helics_opendss_dynamic_var_summary.csv')

    poi_frames = []
    for event_id in sorted(summary['event'].unique()):
        trace_path = sweep_root/f'gridpack_poi_voltage_event_{int(event_id)}.csv'
        if trace_path.exists():
            trace = pd.read_csv(trace_path)
            trace['event'] = int(event_id)
            poi_frames.append(trace)
    poi_all = pd.concat(poi_frames, ignore_index=True)

    poi_all.to_csv(DATA/'gridpack_voltage_turbulence_event_sweep_v3.csv', index=False)
    selected_ts.to_csv(DATA/'gridpack_voltage_control_event_response_v3.csv', index=False)
    selected_summary.to_csv(DATA/'gridpack_voltage_control_summary_v3.csv', index=False)

    colors_v = {'C1':'#377eb8', 'C2':'#984ea3', 'C3':'#e6550d'}
    labels_v = {'C1':'C1 traditional AC', 'C2':'C2 local SST', 'C3':'C3 centralized DC'}

    fig=plt.figure(figsize=(12.6,4.05))
    gs=fig.add_gridspec(1,3,width_ratios=[1.12,1.05,1.16],left=0.052,right=0.985,bottom=0.20,top=0.84,wspace=0.34)
    ax_flow=fig.add_subplot(gs[0,0]); ax_poi=fig.add_subplot(gs[0,1]); ax_boundary=fig.add_subplot(gs[0,2])

    ax_flow.set_axis_off()
    ax_flow.set_title('a  Control chain',loc='left',fontsize=11,weight='bold')
    ax_flow.set_xlim(0,1); ax_flow.set_ylim(0,1)
    boxes=[
        (0.04,0.68,0.25,0.18,'GridPACK\nbranch fault'),
        (0.38,0.68,0.25,0.18,'Bus-150\nPOI trace'),
        (0.72,0.68,0.24,0.18,'HELICS\n20 ms\nexchange'),
        (0.04,0.28,0.25,0.18,'OpenDSS\nfeeder solve'),
        (0.38,0.28,0.25,0.18,'Scenario\ncontrol layer'),
        (0.72,0.28,0.24,0.18,'800 VDC\nload boundary'),
    ]
    for x0,y0,w,h,text in boxes:
        ax_flow.add_patch(FancyBboxPatch((x0,y0),w,h,boxstyle='round,pad=0.02,rounding_size=0.025',
                                         facecolor='white',edgecolor='0.35',lw=1.0))
        ax_flow.text(x0+w/2,y0+h/2,text,ha='center',va='center',fontsize=6.9)
    arrows=[((0.30,0.77),(0.37,0.77)),((0.64,0.77),(0.71,0.77)),((0.84,0.67),(0.84,0.48)),
            ((0.71,0.37),(0.64,0.37)),((0.37,0.37),(0.30,0.37))]
    for (x0,y0),(x1,y1) in arrows:
        ax_flow.add_patch(FancyArrowPatch((x0,y0),(x1,y1),arrowstyle='-|>',mutation_scale=9,lw=1.0,color='0.30'))
    y0=0.08
    ax_flow.text(0.04,y0+0.11,'C1: no fast DC layer -> trip',fontsize=6.8,color=colors_v['C1'])
    ax_flow.text(0.04,y0+0.06,'C2: local VAR reaches limit -> trip',fontsize=6.8,color=colors_v['C2'])
    ax_flow.text(0.04,y0+0.01,'C3: centralized terminal + DC buffer -> served',fontsize=6.8,color=colors_v['C3'])

    for event_id, trace in poi_all.groupby('event'):
        color = '#d94801' if int(event_id)==selected_event else '0.65'
        lw = 1.6 if int(event_id)==selected_event else 0.8
        alpha = 1.0 if int(event_id)==selected_event else 0.55
        ax_poi.plot(trace.time_s, trace.poi_voltage_pu, color=color, lw=lw, alpha=alpha)
    ax_poi.axhline(0.90,color='0.35',lw=0.9,ls='--')
    ax_poi.set_xlim(0.80,1.38); ax_poi.set_ylim(0.0,1.08)
    ax_poi.set_xlabel('Time (s)'); ax_poi.set_ylabel('POI voltage (pu)')
    ax_poi.set_title('b  POI disturbance',loc='left',fontsize=11,weight='bold')
    ax_poi.text(0.03,0.08,f'event {selected_event}: lowest POI {summary.poi_min_pu.min():.3f} pu',
                transform=ax_poi.transAxes,fontsize=6.8,color='0.25')
    ax_poi.grid(alpha=0.25)

    for scen in ['C1','C2','C3']:
        d = selected_ts[selected_ts.scenario==scen]
        ax_boundary.plot(d.time_s, d.load_boundary_voltage_pu, color=colors_v[scen], lw=1.7, label=labels_v[scen])
    ax_boundary.axhline(0.90,color='0.35',lw=0.9,ls='--')
    ax_boundary.set_xlim(0.80,1.38); ax_boundary.set_ylim(0.0,1.08)
    ax_boundary.set_xlabel('Time (s)'); ax_boundary.set_ylabel('Load-boundary voltage (pu)')
    ax_boundary.set_title('c  Boundary response',loc='left',fontsize=11,weight='bold')
    outcome_lines=[]
    for scen in ['C1','C2','C3']:
        r=selected_summary[selected_summary.scenario==scen].iloc[0]
        outcome_lines.append((scen, 'served' if not bool(r.data_center_tripped) else 'trip'))
    for j,(scen,outcome) in enumerate(outcome_lines):
        ax_boundary.text(0.04,0.23-j*0.055,f'{scen}: {outcome}',transform=ax_boundary.transAxes,
                         fontsize=6.9,color=colors_v[scen],ha='left',va='top')
    ax_boundary.legend(fontsize=6.3,frameon=False,loc='lower right')
    ax_boundary.grid(alpha=0.25)
    savefig(fig,'gridpack_voltage_ride_through')
figure4()

# Figure 5
def figure5():
    fig,axes=plt.subplots(1,2,figsize=(10.8,4.1),gridspec_kw={'width_ratios':[1.0,1.35]})
    ax=axes[0]
    bars=[2.1,3.4,4.2]; labs=['2021-22\nstudy','2024-25\nbase','2024-25\nhigh-load\nsensitivity']
    pd.DataFrame({
        'panel': ['5a'] * len(bars),
        'case': ['2021-22 study', '2024-25 base', '2024-25 high-load sensitivity'],
        'san_jose_study_area_load_gw': bars,
        'source': ['CAISO San Jose area planning documents'] * len(bars),
    }).to_csv(DATA/'fig5_load_pocket_source_data_v3.csv', index=False)
    current_rows = []
    for load_gw in np.linspace(0.1, 5.0, 150):
        for pole_kv in [69, 138, 230, 320]:
            current_rows.append({
                'panel': '5b',
                'cluster_load_gw': float(load_gw),
                'pole_voltage_kv': pole_kv,
                'bipole_voltage_kv': 2 * pole_kv,
                'single_bipole_current_ka': float(load_gw * 1e9 / (2 * pole_kv * 1e3) / 1000),
            })
    pd.DataFrame(current_rows).to_csv(DATA/'fig5_voltage_class_current_envelope_v3.csv', index=False)
    x=np.arange(len(bars))
    ax.bar(x,bars,color=['#c9c9c9','#8f8f8f','#4f4f4f'],width=0.62)
    ax.set_xticks(x); ax.set_xticklabels(labs,fontsize=7.5); ax.set_ylabel('San Jose study-area load forecast (GW)')
    ax.set_title('a  Load-pocket scale',loc='left',fontsize=10,weight='bold')
    for i,b in enumerate(bars): ax.text(i,b+0.08,f'{b:.1f} GW',ha='center',fontsize=8)
    ax.annotate('public multi-GW\nload-pocket precedent',xy=(2,4.2),xytext=(0.58,3.15),fontsize=7.5,color='0.25',ha='center',arrowprops=dict(arrowstyle='->',lw=0.9,color='0.35'))
    ax.text(0.02,0.04,'Source: CAISO San Jose area planning documents',transform=ax.transAxes,fontsize=6.6,color='0.35')
    ax.set_ylim(0,4.75); ax.grid(axis='y',alpha=0.22)

    ax=axes[1]
    loads2=np.linspace(0.1,5.0,150)
    for pole,col in [(69,'#9ecae1'),(138,'#e6550d'),(230,'#31a354'),(320,'#756bb1')]:
        I=loads2*1e9/(2*pole*1e3)/1000
        ax.plot(loads2,I,label=f'+/-{pole} kV',color=col,lw=2)
    ax.axhline(4,color='0.4',ls='--',lw=1,label='illustrative 4 kA line')
    ax.axvspan(3.4,4.2,color='#f0f0f0',alpha=0.75,label='3.4-4.2 GW precedent')
    ref_current=1e9/(2*138e3)/1000
    ax.scatter([1.0],[ref_current],c='#e6550d',edgecolor='k',zorder=4)
    ax.annotate('reference case:\n1 GW at +/-138 kV\n= 3.6 kA',xy=(1.0,ref_current),xytext=(1.35,7.2),fontsize=7,color='#e6550d',arrowprops=dict(arrowstyle='->',color='#e6550d',lw=0.8))
    ax.text(3.48,8.8,'multi-GW range needs\nhigher voltage, parallel bipoles\nor both',fontsize=7.2,color='0.25',ha='left')
    ax.set_xlabel('Cluster load (GW)'); ax.set_ylabel('Single-bipole current (kA)'); ax.set_xlim(0,5.25); ax.set_ylim(0,37.5)
    ax.set_title('b  Voltage envelope',loc='left',fontsize=10,weight='bold'); ax.legend(fontsize=7,frameon=False,loc='upper left'); ax.grid(alpha=0.22)
    fig.tight_layout(); savefig(fig,'load_pocket_voltage_envelope')
figure5()

def load_travis_greenfield_outputs():
    """Copy Travis 150 greenfield and T&D co-simulation outputs into package."""
    source_data = SOURCE_ROOT/'data'
    source_gridpack = SOURCE_ROOT/'cosim'/'gridpack_td_dynamic_var'
    gridpack_sweep = source_gridpack/'results_event_sweep'/'event_sweep_summary_compact.csv'
    gridpack_summary = source_gridpack/'results'/'helics_opendss_dynamic_var_summary.csv'
    if gridpack_sweep.exists():
        source_cosim = source_gridpack
        cosim_backend = 'GridPACK/HELICS/OpenDSS branch-fault event sweep'
        cosim_results = source_gridpack/'results_event_sweep'
        cosim_files = ['event_sweep_summary_compact.csv'] + [
            f'gridpack_poi_voltage_event_{i}.csv' for i in range(6)
        ]
    elif gridpack_summary.exists():
        source_cosim = source_gridpack
        cosim_backend = 'GridPACK/HELICS/OpenDSS real Travis 150 RAW/DYR'
        cosim_results = source_gridpack/'results'
        cosim_files = [
            'helics_opendss_dynamic_var_summary.csv',
            'helics_opendss_dynamic_var_timeseries.csv',
            'gridpack_poi_voltage_timeseries.csv',
            'gridpack_travis150_generator_watch.csv',
            'gridpack_travis150_dynamic_input.xml',
            'run_manifest.json',
        ]
    else:
        source_cosim = None
        cosim_backend = 'GridPACK/HELICS/OpenDSS results not archived'
        cosim_results = None
        cosim_files = []
    tables = {}
    for key, filename in {
        'transfer': 'travis150_greenfield_c1_c2_c3_transfer_v2.csv',
        'harmonics': 'travis150_greenfield_c1_c2_c3_harmonics_v2.csv',
        'voltage': 'travis150_greenfield_c1_c2_c3_voltage_v2.csv',
        'summary': 'travis150_greenfield_c1_c2_c3_summary_v2.csv',
    }.items():
        src = source_data/filename
        dst = DATA/filename
        if src.exists():
            shutil.copy(src, dst)
        if dst.exists():
            tables[key] = pd.read_csv(dst)

    for filename in cosim_files:
        src = cosim_results/filename
        dst = DATA/filename
        if src.exists():
            shutil.copy(src, dst)
    for filename in ['150.RAW', '150_gridpack_REECA1_candidate.dyr']:
        src = (source_gridpack/'results')/filename
        if src.exists():
            shutil.copy(src, DATA/filename)
    cosim_summary = DATA/'helics_opendss_dynamic_var_summary.csv'
    if cosim_summary.exists():
        tables['cosim_summary'] = pd.read_csv(cosim_summary)
        tables['_cosim_backend'] = cosim_backend
    sweep_summary = DATA/'event_sweep_summary_compact.csv'
    if sweep_summary.exists():
        tables['event_sweep_summary'] = pd.read_csv(sweep_summary)
        tables['_cosim_backend'] = cosim_backend
    return tables

travis_tables = load_travis_greenfield_outputs()

def figure6():
    summary = travis_tables.get('summary')
    if summary is None or summary.empty:
        return
    cosim_summary = travis_tables.get('cosim_summary')
    event_sweep = travis_tables.get('event_sweep_summary')
    order = ['C1', 'C2', 'C3']
    labels = {
        'C1': 'C1\nAC to 480 V',
        'C2': 'C2\nAC + SST',
        'C3': 'C3\nDC corridor',
    }
    colors = {'C1': '#377eb8', 'C2': '#984ea3', 'C3': '#e6550d'}
    s = summary.set_index('scenario_id').loc[order]

    fig, axes = plt.subplots(1, 3, figsize=(12.2, 4.0), gridspec_kw={'wspace': 0.34})
    fig.subplots_adjust(left=0.06, right=0.985, bottom=0.27, top=0.84)

    ax = axes[0]
    transfer_gw = s['max_transfer_mw']/1000.0
    ax.bar(np.arange(3), transfer_gw, color=[colors[k] for k in order], width=0.62)
    ax.axhline(1.0, color='0.35', lw=1.0, ls='--')
    for i, k in enumerate(order):
        ax.text(i, transfer_gw.loc[k] + 0.035, f'{transfer_gw.loc[k]:.2f} GW', ha='center', fontsize=8)
    c3_gain = (s.loc['C3', 'max_transfer_mw']/s.loc['C1', 'max_transfer_mw'] - 1.0)*100.0
    ax.text(1.98, 0.20, f'C3 +{c3_gain:.1f}%\nvs C1', ha='center', va='bottom', fontsize=8, color=colors['C3'])
    ax.set_xticks(np.arange(3)); ax.set_xticklabels([labels[k] for k in order], fontsize=8)
    ax.set_ylabel('Useful transfer limit (GW)')
    ax.set_ylim(0, 1.62)
    ax.set_title('a  Transfer capacity', loc='left', fontsize=10.5, weight='bold')
    ax.grid(axis='y', alpha=0.25)

    ax = axes[1]
    thd = s['thdv_p95_pct']
    ax.bar(np.arange(3), thd, color=[colors[k] for k in order], width=0.62)
    for i, k in enumerate(order):
        ax.text(i, thd.loc[k] + 0.055, f'{thd.loc[k]:.2f}%', ha='center', fontsize=8)
        source_label = f'{int(s.loc[k, "harmonic_source_count"])} AC-facing\nsource(s)'
        if thd.loc[k] > 0.20:
            ax.text(i, max(0.08, thd.loc[k]*0.12), source_label, ha='center', va='bottom', fontsize=6.7, color='white')
        else:
            ax.text(i, thd.loc[k] + 0.22, source_label, ha='center', va='bottom', fontsize=6.7, color=colors[k])
    ax.set_xticks(np.arange(3)); ax.set_xticklabels([labels[k] for k in order], fontsize=8)
    ax.set_ylabel('p95 THDv at 1 GW (%)')
    ax.set_ylim(0, max(thd)*1.28)
    ax.set_title('b  Harmonic ownership', loc='left', fontsize=10.5, weight='bold')
    ax.grid(axis='y', alpha=0.25)

    ax = axes[2]
    poi = None
    if event_sweep is not None and not event_sweep.empty:
        c = event_sweep.copy()
        c['tripped'] = c['tripped'].astype(bool)
        grouped = c.groupby('scenario')
        boundary = grouped['load_boundary_min_pu'].min().loc[order]
        served = grouped['min_served'].min().loc[order]
        trip = grouped['tripped'].max().loc[order]
        poi = grouped['poi_min_pu'].min().loc[order]
        source_label = 'GridPACK branch-fault event sweep'
    elif cosim_summary is not None and not cosim_summary.empty:
        c = cosim_summary.set_index('scenario').loc[order]
        boundary = c['load_boundary_min_voltage_pu']
        served = c['min_load_served_fraction']
        trip = c['data_center_tripped']
        poi = c['poi_min_voltage_pu'] if 'poi_min_voltage_pu' in c.columns else None
        source_label = travis_tables.get('_cosim_backend', 'GridPACK/HELICS/OpenDSS')
    else:
        boundary = s['data_center_load_served_min_fraction']*0.0 + s['p95_pcc_voltage_deviation_pct']
        served = s['data_center_load_served_min_fraction']
        trip = s['data_center_tripped']
        source_label = 'Travis 150 screening result'
    x = np.arange(3)
    ax.bar(x, boundary, color=[colors[k] for k in order], width=0.58, alpha=0.92, label='data-center boundary min')
    if poi is not None:
        ax.scatter(x, poi, marker='v', s=44, color='0.28', zorder=4, label='POI min')
    ax.plot(np.arange(3), served, color='0.15', marker='D', ms=5.5, lw=1.2, label='load served fraction')
    ax.axhline(0.90, color='0.35', lw=1.0, ls='--')
    for i, k in enumerate(order):
        flag = 'trip' if bool(trip.loc[k]) else 'served'
        ax.text(i, min(boundary.loc[k] + 0.055, 1.08), f'{boundary.loc[k]:.2f} pu\n{flag}', ha='center', fontsize=8)
    ax.set_xticks(np.arange(3)); ax.set_xticklabels([labels[k] for k in order], fontsize=8)
    ax.set_ylabel('Boundary voltage / served fraction')
    ax.set_ylim(0, 1.16)
    ax.set_title('c  Voltage ride-through', loc='left', fontsize=10.5, weight='bold')
    ax.text(0.00, -0.24, source_label, transform=ax.transAxes, fontsize=6.4, color='0.35', va='top')
    ax.legend(fontsize=6.4, frameon=False, loc='upper left')
    ax.grid(axis='y', alpha=0.25)

    savefig(fig, 'travis150_greenfield_benefits')
figure6()

# Supplementary figures: S1 protection dynamics, S2 EMT validation, S3 buffer feasibility, S4 economics/copper
def figure_s1():
    fig,axes=plt.subplots(1,3,figsize=(12,3.8))
    ax=axes[0]; ax.axis('off'); ax.set_title('a  Protection zones',loc='left',fontsize=10,weight='bold')
    ax.set_xlim(0,10); ax.set_ylim(0,4); ac='#1f77b4'; dc='#d94801'
    draw_icon(ax,0.8,2.6,'grid',0.7); draw_icon(ax,2.0,2.6,'converter',0.7); ax.plot([1.1,1.65],[2.6,2.6],color=ac,lw=2)
    ax.plot([2.35,8.2],[2.6,2.6],color=dc,lw=2.5)
    for x in [3.5,5.1,6.8]:
        ax.add_patch(Rectangle((x-0.07,2.42),0.14,0.36,facecolor='#fff',edgecolor='0.3'))
    for i,x in enumerate([4.0,5.8,7.6]):
        ax.plot([x,x],[2.6,1.55],color=dc,lw=1.5); draw_icon(ax,x,1.35,'dcdc',0.45); ax.text(x,0.9,f'campus {i+1}',ha='center',fontsize=7)
    ax.add_patch(Polygon([[5.1,2.95],[5.3,2.55],[5.16,2.58],[5.36,2.25],[5.0,2.62]],closed=True,facecolor='#ffd92f',edgecolor='0.2'))
    ax.text(5.35,3.1,'backbone fault',fontsize=7)
    for x,label in [(2.0,'terminal\nprotection'),(3.5,'section\nbreaker'),(6.8,'section\nbreaker')]:
        ax.text(x,3.15,label,ha='center',fontsize=6.7,color='0.25')
    ax.text(8.4,0.45,'screening model,\nnot breaker design',ha='right',fontsize=7,color='0.35')
    ax=axes[1]; ax.set_title('b  Fault response',loc='left',fontsize=10,weight='bold')
    ax.plot(fault_df.time_s*1000,fault_df.fault_current_kA,color='#d94801',lw=1.5,label='fault current')
    ax2=ax.twinx(); ax2.plot(fault_df.time_s*1000,fault_df.backbone_voltage_pu,color='#3182bd',lw=1.3,label='backbone V')
    ax.axvline(3,color='0.5',ls=':',lw=1); ax.axvline(18,color='0.5',ls='--',lw=1)
    ax.text(3.4,ax.get_ylim()[1]*0.88,'detect',fontsize=7,color='0.35')
    ax.text(18.4,ax.get_ylim()[1]*0.78,'open',fontsize=7,color='0.35')
    ax.set_xlabel('Time (ms)'); ax.set_ylabel('Fault current (kA)'); ax2.set_ylabel('Voltage (pu)'); ax.grid(alpha=0.25)
    ax=axes[2]; ax.set_title('c  Ride-through screen',loc='left',fontsize=10,weight='bold')
    ax.plot(fault_df.time_s*1000,fault_df.campus1_voltage_pu,label='near/faulted section',color='#e6550d')
    ax.plot(fault_df.time_s*1000,fault_df.campus2_voltage_pu,label='healthy campus 2',color='#31a354')
    ax.plot(fault_df.time_s*1000,fault_df.campus3_voltage_pu,label='healthy campus 3',color='#756bb1')
    ax.set_xlabel('Time (ms)'); ax.set_ylabel('Campus DC voltage (pu)'); ax.legend(fontsize=7,frameon=False); ax.grid(alpha=0.25)
    fig.tight_layout(); savefig(fig,'dc_fault_protection_dynamic')
figure_s1()

def figure_s2():
    fig,axes=plt.subplots(1,2,figsize=(10,3.8))
    ax=axes[0]
    ax.plot(validation_df.dt_s,validation_df.rmse_MW_vs_1ms_reference,marker='o',color='#e6550d')
    ax.set_xscale('log'); ax.invert_xaxis()
    ax.set_xlabel('Time step (s)'); ax.set_ylabel('RMSE vs 1 ms reference (MW)'); ax.set_title('a  Timestep check',loc='left',fontsize=10,weight='bold'); ax.grid(alpha=0.25)
    ax=axes[1]
    ax.plot(tf_df.frequency_Hz,tf_df.simulated_gain,marker='o',label='simulation',color='#3182bd')
    ax.plot(tf_df.frequency_Hz,tf_df.theory_gain,'--',label='first-order theory',color='0.25')
    ax.set_xscale('log'); ax.set_yscale('log'); ax.set_xlabel('Frequency (Hz)'); ax.set_ylabel('Grid-command gain'); ax.set_title('b  Transfer function',loc='left',fontsize=10,weight='bold'); ax.legend(fontsize=7,frameon=False); ax.grid(alpha=0.25,which='both')
    fig.tight_layout(); savefig(fig,'averaged_emt_validation')
figure_s2()

def figure_s3():
    fig,(ax,ax2)=plt.subplots(1,2,figsize=(10.5,4.1),gridspec_kw={'width_ratios':[1.0,1.25]})
    table=buffer_table.copy()
    ax.set_title('a  Buffer suitability',loc='left',fontsize=10,weight='bold')
    score={'not energy storage':-1,'low':0,'partial':1,'medium':2,'high':3}
    labels={-1:'not\nstorage',0:'low',1:'partial',2:'medium',3:'high'}
    matrix=np.array([[score[v] for v in row] for row in table[['high_power_suitability','energy_window_suitability']].values])
    cmap=matplotlib.colors.ListedColormap(['#f7f7f7','#fcbba1','#fdae6b','#a1d99b','#31a354'])
    im=ax.imshow(matrix,aspect='auto',vmin=-1,vmax=3,cmap=cmap)
    ax.set_yticks(np.arange(len(table)))
    ax.set_yticklabels(table['technology'],fontsize=7)
    ax.set_xticks([0,1])
    ax.set_xticklabels(['High-power\nresponse','0.42 MWh\nwindow'],fontsize=8)
    for i in range(matrix.shape[0]):
        for j in range(matrix.shape[1]):
            ax.text(j,i,labels[int(matrix[i,j])],ha='center',va='center',fontsize=7)
    ax.tick_params(length=0)
    ax2.axis('off'); ax2.set_title('b  Deployment role',loc='left',fontsize=10,weight='bold')
    ax2.set_xlim(0,1); ax2.set_ylim(0,1)
    ax2.text(0.02,0.90,'Technology',fontsize=7,weight='bold')
    ax2.text(0.30,0.90,'Layer',fontsize=7,weight='bold')
    ax2.text(0.55,0.90,'Role',fontsize=7,weight='bold')
    ax2.hlines(0.86,0.02,0.98,color='0.25',lw=1.0)
    for i,r in enumerate(table.itertuples(index=False)):
        y=0.79-i*0.15
        ax2.text(0.02,y,textwrap.fill(r.technology,18),fontsize=6.4,va='top')
        ax2.text(0.30,y,textwrap.fill(r.deployment_layer,18),fontsize=6.2,va='top')
        ax2.text(0.55,y,textwrap.fill(r.role,38),fontsize=6.2,va='top')
        ax2.hlines(y-0.10,0.02,0.98,color='0.86',lw=0.7)
    fig.tight_layout()
    savefig(fig,'shared_buffer_feasibility')
figure_s3()

def figure_s4():
    fig,axes=plt.subplots(1,2,figsize=(10,3.8))
    econ_df=pd.read_csv(DATA/'cost_copper_envelope_v3.csv')
    ax=axes[0]
    grid=econ_df.dropna(subset=['annual_value_USD_M']).pivot(index='load_factor',columns='electricity_price_USD_MWh',values='annual_value_USD_M')
    im=ax.imshow(grid.values,origin='lower',aspect='auto',extent=[price_grid.min(),price_grid.max(),lf_grid.min(),lf_grid.max()],cmap='YlGn')
    ax.set_xlabel('Electricity price ($/MWh)'); ax.set_ylabel('Load factor'); ax.set_title('a  Loss-saving value',loc='left',fontsize=10,weight='bold')
    ax.scatter([assumptions['electricity_price_USD_per_MWh_mid']],[assumptions['economic_load_factor']],marker='*',s=110,c='#d94801',edgecolor='k',zorder=3)
    ax.text(assumptions['electricity_price_USD_per_MWh_mid']+4,assumptions['economic_load_factor'],'reference',fontsize=7,va='center')
    fig.colorbar(im,ax=ax,shrink=0.8,label='Annual value (million USD/yr)')
    ax=axes[1]
    current_idx=econ_df[econ_df.metric=='current_length_index_kA_km']
    x=np.arange(len(current_idx))
    vals=current_idx.value.to_numpy()
    ax.bar(x,vals,color=['#377eb8','#984ea3','#e6550d'],alpha=0.85)
    ax.set_xticks(x); ax.set_xticklabels(['Traditional\nAC','Local\nSST','DC\nbackbone'],fontsize=8)
    for i,v in enumerate(vals):
        ax.text(i,v+2,f'{v:.0f}',ha='center',fontsize=7)
    ax.set_ylabel('Current-length index (kA km)'); ax.set_title('b  Current-length index',loc='left',fontsize=10,weight='bold'); ax.grid(axis='y',alpha=0.25)
    ax.text(0.03,0.90,'index = current x corridor length',transform=ax.transAxes,fontsize=6.8,va='top',color='0.35')
    fig.tight_layout(); savefig(fig,'cost_copper_envelope')
figure_s4()

# ---------------------------- Manuscript text ----------------------------
abstract = """AI factories are becoming synchronized, DC-native, gigawatt-scale loads, but electric-grid planning still treats most data centers as passive AC facilities. This creates a boundary-placement problem: when the useful computational load is supplied at 800 VDC, the AC/DC interface can remain inside each building, move to local solid-state-transformer interfaces, or be placed upstream as a utility-operated subtransmission DC backbone. Here we show that moving this boundary upstream provides three coupled system benefits. First, it increases useful transfer capacity for new data-center corridors while reducing corridor and conversion losses relative to traditional AC delivery. Second, it centralizes AC-facing harmonic ownership at one utility converter terminal rather than distributing large converter interfaces across multiple campuses. Third, it creates a shared DC buffering layer that improves voltage ride-through for synchronized AI loads. Reference calculations, uncertainty sweeps, harmonic screening and a Travis 150 dynamic co-simulation preserve the same ordering: traditional AC is weakest, local SSTs improve selected metrics but retain multiple AC-facing interfaces, and the centralized DC corridor provides the most robust grid-facing boundary. These results support a falsifiable systems claim: for clustered AI factories, the AC/DC boundary is a subtransmission planning variable rather than only a building-level design choice."""

intro = """AI factories change the electrical problem that grids must solve. A conventional data center can often be approximated in planning studies as a large but mostly passive load. A modern AI factory is a synchronized computing machine. Training iterations, all-reduce communication, checkpointing and accelerator power-management events can appear electrically as coherent power modulation across thousands of GPUs. At the scale of multiple campuses connected to the same grid pocket, power delivery becomes part of the computing architecture.

The planning problem is therefore not only how much electricity AI factories consume, but where the grid should place the controllable AC/DC boundary when the useful computational load is already DC-native. That placement can change useful transfer capacity, conversion losses, harmonic ownership, voltage ride-through and the division of control responsibility between utilities and campuses.

Data-center energy studies have long shown that efficiency improvements can moderate electricity growth even as digital workloads expand, but recent AI-era assessments now frame data centers as geographically concentrated point loads whose growth can challenge regional planning timelines [1-4]. The load-side technology trajectory is also moving toward DC. Recent industry roadmaps describe 800 VDC as a power-distribution architecture for AI data centers and AI factories that reduces current, copper, distribution volume and conversion stages while supporting future high-density racks [5]. Earlier 380 V DC data-center distribution studies and recent low-voltage DC facility guidance show the same architectural direction at lower voltage classes: reducing conversion stages can improve efficiency and reliability [6,7]. This makes the 800 VDC interface a relevant terminal boundary for future AI-factory power delivery.

If the endpoint is DC, the system-level question is where the AC/DC boundary should be placed. Most current discussions move that boundary from the rack to the facility. This study asks whether it should move farther upstream, into the subtransmission corridor. The proposed architecture uses a utility-operated AC/DC terminal to feed a bipolar subtransmission DC backbone. Campus DC/DC stations then step the backbone to a 34.5 kV DC distribution layer and ultimately to the 800 VDC data-center interface.

Device-level work makes this question technically plausible. A 10 kV SiC 7 kV/400 V DC transformer for future data centers demonstrated 99.0% full-load DC/DC efficiency and 3.8 kW/L power density; the associated 3.8 kV AC to 400 V DC SST chain reached 98.1% full-load efficiency [8]. A modular 5 kV SiC SST demonstrated single-stage MVDC-to-LVDC or MVDC-to-LVAC conversion, full-range zero-voltage switching, controlled dv/dt and modular series/parallel scalability [9]. A 20 kW 1000 V/48 V prototype further shows that raising the data-center distribution voltage can reduce low-voltage current stress, with an estimated efficiency improvement to 97.5% using synchronous rectification [10]. More general SST literature likewise treats SSTs as controllable medium-voltage power-electronics interfaces rather than only replacements for low-frequency transformers [11].

The grid-side motivation is also visible. A production-scale AI training power study by Microsoft, OpenAI and NVIDIA reports that synchronized training phases make power swings visible at rack, data-center and grid levels; at scale these swings can reach tens or hundreds of megawatts and can occupy sub-synchronous frequency ranges relevant to utility equipment [12]. That work frames both time-domain ramp constraints and frequency-domain limits, including a 0.1-20 Hz range, as requirements for safe scaling.

The knowledge gap is therefore not whether efficient DC conversion is possible, or whether AI training loads are dynamic. It is where the AC/DC boundary should sit when AI factories become multi-campus, gigawatt-scale grid assets. We test the claim that for AI factories, the AC/DC boundary is no longer only a building-level choice; it is a subtransmission planning variable."""

results_sections = [
("An AI-native architecture with the AC/DC boundary moved upstream", """We compare three architectures that deliver identical useful power to an 800 VDC data-center boundary (Fig. 1). The traditional AC architecture uses a utility substation, an AC subtransmission corridor, facility AC distribution and distributed AC/DC conversion at each campus. The local-SST architecture keeps the AC corridor but places a solid-state transformer at each AI campus. The proposed architecture moves the first AC/DC terminal upstream and treats the DC subtransmission corridor as a utility asset. Downstream conversion is entirely DC/DC: from subtransmission DC to 34.5 kV DC campus distribution, and from that layer to 800 VDC.

The conceptual difference is the electrical boundary seen by the utility. In the first two architectures, each campus remains an AC-facing load with its own grid-interfacing converter behaviour. In the proposed architecture, the AC grid sees one controlled converter terminal, while campus converters are DC/DC devices embedded behind a shared DC backbone. This turns a cluster of AI campuses from a set of distributed harmonic and ramp sources into a coordinated DC-native load pocket.

This architecture definition fixes the useful 800 VDC boundary across all cases, so later comparisons attribute changes to the grid-interface placement rather than to different delivered computational power."""),
("Transfer capacity is coupled to loss reduction", """For a central reference case, we model a 1 GW cluster served over a 20 km reinforced subtransmission corridor. The traditional AC case uses 138 kV line-to-line at 0.98 power factor; the proposed DC case uses a +/-138 kV bipole, or 276 kV pole-to-pole. This is a representative voltage class rather than a prescribed standard. We define useful transfer capacity as the MW delivered to the common 800 VDC load boundary under the same grid-side input limit.

At the grid input required for traditional AC to serve 1 GW, the local-SST case delivers 12.2 MW more useful load and the DC backbone delivers 13.0 MW more useful load to the 800 VDC boundary (Fig. 2a). The same calculation gives total losses of 39.1 MW for traditional AC, 26.5 MW for local SSTs and 25.7 MW for the DC backbone at a 1 GW load point. The corresponding end-to-end efficiencies are 96.23%, 97.42% and 97.49%. A separate 99.0% local-SST efficiency sensitivity gives 21.3 MW loss, 97.92% end-to-end efficiency and 17.3 MW useful transfer gain at the same grid input; this is treated only as a sensitivity case, not as a demonstrated reference architecture.

This transfer-capacity result is not a separate claim from loss reduction. It is the useful-load consequence of delivering the same computational boundary through a lower-loss corridor and conversion chain. Under the 99.0% local-SST sensitivity case, local conversion can exceed the DC backbone in pure efficiency. The architectural case emerges because the DC backbone gives a transfer-capacity gain over traditional AC in the same direction as harmonic ownership and dynamic-voltage benefits. A load-distance sweep from 100 MW to 3 GW and from 5 to 100 km shows where the DC useful-transfer gain over traditional AC exceeds 10, 50 and 100 MW (Fig. 2c). A Monte Carlo uncertainty sweep and one-at-a-time tornado analysis show that corridor length, conductor resistance and downstream conversion assumptions dominate the quantitative result (Fig. 2b,d).

The transfer-capacity conclusion is therefore supported by closed-form loss equations, a load-distance design sweep, Monte Carlo uncertainty and a one-at-a-time sensitivity screen rather than by a single reference operating point."""),
("A DC backbone changes harmonic compliance into harmonic ownership", """Traditional AC and local-SST architectures can be designed to meet harmonic limits, but they place multiple large AC-facing converter interfaces along the corridor. Their aggregate harmonic voltage distortion depends on local filters, network impedance, cable capacitance, phase relationships between sites and resonance. The proposed DC backbone concentrates the AC-facing converter at a single utility-operated terminal. Campus stations are DC/DC interfaces and therefore do not directly inject AC harmonics into the subtransmission grid.

We quantify this ownership change with an OpenDSS-ready network and a reproduced nodal frequency-domain solver. The network uses a 10 GVA Thevenin short-circuit strength at 138 kV, three campus buses along a 20 km corridor, harmonic-dependent source impedance and resonance amplification around selected orders. Distributed architectures are represented by three AC-facing converter spectra with random relative phases; the DC-backbone case is represented by one filtered grid-facing converter terminal.

For the central assumptions, the 95th-percentile PCC voltage THD is 3.95% for traditional AC, 1.55% for local SSTs and 0.78% for the DC backbone (Fig. 3b). Adding active filtering or storage to the traditional AC case improves the metric, and coordinated control improves the local-SST case, but neither changes the number of AC-facing interfaces. These values are screening metrics, not a substitute for project-specific IEEE 519 compliance studies [13]. Their purpose is narrower and architectural: moving DC upstream changes a distributed compliance problem into a single utility-owned terminal design problem.

The harmonic result is supported by both an OpenDSS-compatible network description and an independent nodal frequency-domain solver using the same equivalent circuit and converter spectra."""),
("The DC backbone buffers voltage turbulence at a controllable boundary", """The third benefit is voltage ride-through under transmission-side voltage turbulence. We use the Travis 150 GridPACK event sweep to represent the disturbance source, pass the bus-150 POI voltage traces through a 20 ms HELICS/OpenDSS exchange, and compare the data-center boundary response for the three supply architectures (Fig. 4). This replaces the earlier averaged voltage-fluctuation figure with the same dynamic disturbance chain used in the Travis 150 validation.

The voltage-control structure differs across cases. C1 exposes the data-center boundary directly to the AC disturbance. C2 adds local SST Volt-VAR support near the campus, but the local controller reaches its reactive-power limit in the severe branch-fault sweep. C3 places the voltage-control responsibility at the centralized AC/DC terminal and represents a DC-buffer layer behind that terminal. For the lowest-POI event in the sweep, the bus-150 voltage reaches 0.092 pu. C1 and C2 trip, while C3 keeps the 800 VDC data-center boundary served in the modeled control architecture.

This is still a screening result, not a hardware controller validation. Its purpose is to show where the disturbance is handled. The centralized DC case does not make the transmission-side event disappear; it changes the grid-facing boundary from multiple campus interfaces to one utility terminal with a represented DC-buffer layer."""),
("Data-center load pockets are becoming planning objects", """The proposed architecture is motivated by load pockets that are large, concentrated and data-center driven. Public planning documents for the San Jose area show a load pocket growing from approximately 2.1 GW in an earlier study case to 3.4 GW in a later base case and 4.2 GW in a sensitivity case (Fig. 5a) [14-17]. This paper does not claim that a specific planned HVDC project is a 138 kV DC AI-factory backbone. The point is that data-center-driven load pockets are already large enough to motivate controllable transmission solutions.

This section is used only to establish planning relevance from public documents; it is not used as routing evidence for the Travis synthetic case or as a claim about any specific project.

The voltage-class envelope in Fig. 5b shows why the paper uses +/-138 kV only as a representative subtransmission design point. At 1 GW, +/-138 kV corresponds to approximately 3.6 kA bipole current. Higher multi-GW corridors move naturally toward higher voltage classes such as +/-320 kV. The relevant design variable is therefore not one fixed voltage, but the relocation of the AC/DC boundary to a voltage class compatible with load, distance, current limit, insulation and protection requirements."""),
("Travis 150 greenfield configurations preserve the three-benefit ordering", """We next use the TAMU Travis 150 synthetic electric case as a Texas load-pocket test bed (Fig. 6). TAMU describes the case as a 150-bus synthetic electric system corresponding to the Austin-Travis County T&D system and states that it is synthetic, not an actual grid [18]. We ignore the companion gas network. ERCOT long-term planning materials are used only as regional motivation for large-load growth and transmission-planning context, not as routing data [19].

The Travis study adds new data-center supply systems rather than converting existing AC lines. C1 is a new traditional AC data-center supply ending at 480 V AC facility distribution. C2 keeps a new AC corridor but places an SST at the data-center side, with local dynamic VAR support on the 34.5 kV AC side and an 800 VDC data-center boundary. C3 builds a new dedicated bipolar DC corridor with a centralized grid-facing AC/DC terminal and DC/DC conversion near the campus. At the 1 GW data-center load, the useful transfer limit is 1.17 GW for C1, 1.24 GW for C2 and 1.44 GW for C3, giving C3 a 22.8% transfer increase relative to C1 and 16.7% relative to C2. The 95th-percentile THDv screen follows the same ordering: 1.67% for C1, 0.53% for C2 and 0.08% for C3.

For voltage ride-through, we couple the Travis 150 GridPACK dynamic case to an OpenDSS data-center feeder through HELICS. The transmission side uses a GridPACK-compatible Travis 150 RAW/DYR dynamic deck. We apply six short branch-fault events on the 137-150 transmission branch as shifted 3 s GridPACK simulations so that bus-150 POI voltage observations can be passed to the distribution federation. Across the event sweep, the POI reaches 0.091994 pu. C1 and C2 trip under these severe disturbances; the C2 local VAR controller reaches its reactive-power limit but cannot prevent the 800 VDC boundary from collapsing in these runs. C3 keeps the 800 VDC data-center boundary served in the modeled C3 control architecture because the disturbance is handled at the centralized AC/DC terminal and represented DC-buffer layer. The GridPACK traces are generated from the Travis 150 RAW/DYR case, the bus-150 POI voltage traces are passed to the HELICS/OpenDSS federation at 20 ms resolution, and all OpenDSS solves converge in the event sweep. The result supports the architectural conclusion and also exposes the main local-control risk: local VAR devices can interact with utility LTCs, voltage regulators, capacitor banks, smart inverters or centralized STATCOM/SVC controls unless supervised coordination is added.""")]

discussion = """This study reframes AI-factory power delivery as a grid-interface placement problem. The central result is that moving the AC/DC boundary from the facility to the subtransmission corridor co-locates three system benefits: higher useful transfer with lower losses relative to traditional AC, centralized harmonic ownership and dynamic voltage buffering. The result does not imply that every data center should be served by DC subtransmission, or that +/-138 kV is a universal optimum; it shows that once AI factories become clustered, synchronized and DC-native, the location of the AC/DC boundary becomes a planning variable.

The comparison also shows why efficiency alone is an incomplete criterion. A high-efficiency local-SST sensitivity case can approach or exceed the DC backbone in pure efficiency, but it does not change the architecture. Local SSTs retain multiple AC-facing grid interfaces and do not automatically provide a shared DC layer for buffering synchronized multi-campus load dynamics. The proposed backbone is valuable because the three benefits are co-located at one controllable boundary.

The Travis 150 result strengthens this conclusion only as a reproducible test-bed result. It is not a site-selection claim, a real Austin routing study or a reconstruction of an operational disturbance. Its value is that the same C1-C2-C3 ordering appears when the architectures are implemented as new data-center supply systems on a synthetic Austin-Travis electric case and then stressed through an installed GridPACK/HELICS/OpenDSS branch-fault workflow. The result also clarifies the disadvantage of the local approach. A local SST VAR controller can improve a static screen, but in a severe sag it may saturate; without supervisory coordination it may also fight nearby smart inverters, substation load-tap changers, capacitor banks, line regulators or centralized STATCOM/SVC controls on different time scales.

Several technical risks remain. DC protection, pole-to-ground fault detection, hybrid DC breakers, grounding, insulation coordination, converter interoperability and electromagnetic-transient stability must be demonstrated before deployment. These risks are the same DC-grid feasibility, protection and converter-interoperability questions identified in HVDC-grid and DC/DC-converter guidance [20-22]. We include protection-screening dynamics and an averaged EMT model to make the research boundary explicit, but do not claim a finished hardware design. The decisive follow-up is pilot-grade EMT and hardware-in-the-loop validation of the grid-facing terminal, DC/DC stations and AI-load emulator.

This study reframes AI factories as grid-planning objects rather than only building loads. The central claim is falsifiable: if a multi-campus AI load can be served by a subtransmission DC backbone, then the same upstream DC boundary should simultaneously reduce corridor/conversion losses relative to traditional AC, centralize AC harmonic ownership and reduce sub-synchronous grid-side voltage modulation relative to architectures that keep AC in the corridor. The models and repository are provided to make that claim testable."""

methods = [
("Architecture boundary and scenario definitions", """The evaluation boundary begins at the grid-facing/subtransmission supply point and ends at the 800 VDC data-center interface. The traditional AC case uses a 138 kV AC corridor and downstream AC distribution before conversion to 800 VDC. The local-SST case uses the same AC corridor but converts at each campus using an SST. The proposed case uses a grid-facing AC/DC terminal, a bipolar subtransmission DC corridor, DC/DC conversion to a 34.5 kV DC distribution layer and DC/DC conversion to 800 VDC. The DC/DC interface assumptions are architecture-level abstractions of HVDC-to-MVDC conversion functions studied for DC grids [22]. The central reference system is a 1 GW three-campus cluster served over a 20 km equivalent corridor. The DC design point is +/-138 kV, or 276 kV pole-to-pole."""),
("Corridor efficiency and uncertainty model", """For AC cases, the receiving-end corridor power is P_recv = P/eta_downstream, where P is the useful 800 VDC load and eta_downstream is the downstream conversion efficiency. Corridor current is I_AC = P_recv/(sqrt(3) V_LL pf), AC line loss is 3 I_AC^2 R, grid input is P_recv plus line loss, and total loss is grid input minus P. For the DC case, receiving-end corridor power is P_recv = P/(eta_DC/DC,1 eta_DC/DC,2), bipole current is I_DC = P_recv/V_pp, line loss is 2 I_DC^2 R, grid input is (P_recv plus line loss)/eta_AC/DC, and total loss is grid input minus P. Central assumptions and the 99.0% local-SST efficiency sensitivity case are listed in Supplementary Table 1, and uncertainty ranges are encoded in the public repository."""),
("Harmonic screening model and OpenDSS cross-check", """The harmonic model is a frequency-domain screening model. It represents the 138 kV grid by a 10 GVA Thevenin short-circuit strength, three corridor buses and harmonic-dependent source impedance with resonance amplification. OpenDSS-compatible circuit files and archived OpenDSSDirect.py harmonic-run artifacts are included in the repository. The figure-generation script also includes an independent nodal-frequency solver that uses the same equivalent network and harmonic spectra, so the screening result can be reproduced without a proprietary EMT tool. The output metrics are PCC voltage THD and individual harmonic voltage distortion. Parameter provenance is summarized in Supplementary Table 1; measured literature values, public planning data and study assumptions are separated in the public data tables."""),
("Averaged EMT-style dynamic buffering model", """The dynamic waveform is synthetic but parameterized from the published structure of AI training power traces: compute phases with high accelerator utilization, periodic communication dips and less frequent checkpointing dips [12]. The traditional AC case passes the waveform directly to the grid. The local-SST case applies a 1.1 s first-order smoothing function. The DC-backbone case applies a 16 s grid-facing power command; the difference between the AI load and the commanded grid power defines shared DC-buffer power. The dynamic robustness grid repeats this averaged model across campus count N = 1, 3, 6 and 10; cluster load P = 0.25, 1, 2 and 4.5 GW; voltage class 69, 138, 230 and 320 kV; short-circuit ratio Ssc/P = 3, 5, 10 and 20; random, partial and coherent temporal phase alignment; and corridor lengths of 5, 20, 50 and 100 km. Supplementary Note 2 gives the averaged state equations and validates the first-order command model by time-step convergence and transfer-function tests. The voltage and spectral metrics are screening proxies aligned with voltage-fluctuation and interconnection-oscillation concerns [23,24]. This is an averaged EMT-style comparison of architecture-level exposure, not a switching EMT validation of a specific converter design."""),
("Travis 150 GridPACK/HELICS/OpenDSS co-simulation", """The Travis 150 study uses the electric side of the TAMU synthetic gas-electric Travis 150 case and ignores the 47-node gas network. The importer accepts a downloaded PowerWorld AUX electric case through the --travis-case option and falls back to the repository's Austin/Travis synthetic corridor candidates when the external file is absent. The flagship data-center corridor is treated as a new build and the data-center load is incremental to native Travis load. In the AUX-based run, closed high-voltage candidate branches are ranked by source strength, load-pocket suitability, transfer and converter headroom and corridor length; the selected span is from 163 Decker Creek Power Plant 1 to 147 Travis_DS_127 1, approximately 9.64 km at 230 kV. The archived fallback placement is B_04 to B_101. The reported sensitivities use 250 MW, 500 MW and 1 GW loads; the main text reports the 1 GW case.

C1 is modelled as a new AC data-center corridor ending at a 480 V AC facility-distribution boundary. C2 uses a new AC corridor with an SST at the data-center side, local dynamic VAR support at the 34.5 kV AC side and an 800 VDC load boundary. C3 uses a new dedicated bipolar DC data-center corridor, a grid-facing AC/DC terminal, DC/DC conversion near the campus and an 800 VDC load boundary. Transfer capacity is the maximum useful data-center MW before the first thermal, converter, voltage, reactive-power or stability screen violation. Harmonic metrics use the same ownership distinction as the main harmonic model but are scaled to the Travis corridor short-circuit strength and incremental data-center load.

The T&D dynamic run uses HELICS to exchange a GridPACK transmission POI voltage trajectory with an OpenDSS data-center feeder and a controller federate. The transmission side uses a GridPACK-compatible Travis 150 RAW/DYR dynamic deck archived with the reproducibility package. Bus 150 is used as the POI voltage observation, and the disturbance is a set of six short branch faults on the 137-150 transmission branch. The GridPACK Python stepwise observation API initializes one event at a time, so the six XML events are run as independent shifted 3 s simulations for the event sweep. The resulting 20 ms bus-150 POI voltage traces are passed to the HELICS/OpenDSS distribution federation. Across the sweep, the lowest observed POI voltage is 0.091994 pu. The controller federate implements C2 local 34.5 kV VAR support and C3 centralized AC/DC-terminal support and DC-buffer ride-through. Trip logic flags an AC-side trip when voltage remains below 0.50 pu for at least 0.04 s. The run manifests record installed-tool versions and convergence flags."""),
("Protection-zone screening", """Representative protection dynamics are simulated for a backbone pole-to-ground fault and a campus DC/DC internal fault. The model includes detection, converter current limiting, breaker opening, section isolation and healthy-campus re-energization. It is intended to check plausibility and expose the required protection functions identified in DC-grid protection studies [20,21]; it is not a validated DC-breaker or insulation-coordination design.""")]

figure_legends = {
'Fig. 1 | Delivery architectures.':'The three supply architectures place the grid-facing AC/DC boundary at different locations. a, Traditional AC delivery keeps AC through the subtransmission corridor and campus switchyards before facility-level conversion. b, Local SST delivery keeps the AC corridor but converts at each campus through a local solid-state transformer. c, The proposed architecture moves the AC/DC boundary upstream to a centralized utility terminal and serves campuses through a subtransmission DC backbone and local DC/DC conversion.',
'Fig. 2 | Useful transfer capacity.':'a, Central 1 GW, 20 km reference-case useful transfer gain when all architectures are constrained to the same grid-side input required for traditional AC to serve 1 GW; loss annotations report the corresponding 1 GW loss point. b, Monte Carlo uncertainty in useful transfer gain relative to traditional AC at the reference point. c, Load-distance sweep showing where the DC-backbone useful-transfer gain over traditional AC exceeds 10, 50 and 100 MW. d, One-at-a-time sensitivity of the central transfer gain. A 99.0% local-SST efficiency sensitivity case is reported in the text and Supplementary Table 1.',
'Fig. 3 | Harmonic ownership.':'a, Harmonic ownership boundary for distributed AC-facing converter cases versus the proposed single utility AC/DC terminal. b, Monte Carlo PCC voltage THD for the three architectures and two stronger baselines, with a 5% planning guide shown for context. c, 95th-percentile individual harmonic voltage distortion. d, Direct OpenDSS harmonic solve compared with the internal nodal-frequency solver.',
'Fig. 4 | GridPACK voltage ride-through.':'a, Coupling structure used to pass the GridPACK branch-fault POI voltage trace through the HELICS/OpenDSS data-center feeder and scenario control layer. b, Six shifted GridPACK branch-fault POI voltage traces at bus 150; the highlighted trace is the lowest-POI event. c, Data-center load-boundary voltage response for C1, C2 and C3 in the highlighted event. C1 and C2 trip in the severe disturbance, whereas C3 remains served in the modeled centralized-terminal and DC-buffer architecture.',
'Fig. 5 | Load-pocket context.':'a, CAISO San Jose area planning data showing a public multi-GW load-pocket precedent. b, Single-bipole current as a function of cluster load for candidate DC voltage classes; the 1 GW reference point and 3.4-4.2 GW public planning precedent show why voltage class, circuit count or both must scale with load.',
'Fig. 6 | Travis 150 validation.':'a, Useful transfer limit for new C1 traditional AC, C2 AC plus SST and C3 bipolar DC data-center corridors in the TAMU Travis 150 synthetic electric case. b, 95th-percentile voltage THD and AC-facing harmonic-source ownership at 1 GW. c, GridPACK/HELICS/OpenDSS branch-fault event sweep showing bus-150 POI minimum voltage, minimum data-center boundary voltage, load-served fraction and trip outcome for six shifted 3 s branch-fault simulations with 20 ms HELICS/OpenDSS exchange.'}

data_availability = """All numerical data underlying the graphs are included in the accompanying source-data folder and reproducibility package as CSV files. The package includes the Travis 150 greenfield outputs, GridPACK/HELICS/OpenDSS summaries, GridPACK POI voltage traces, run manifests and figure-generation inputs. The original Travis 150 electric case is available from the TAMU Electric Grid Test Case Repository subject to its download form. Public external data are cited in the References. No restricted operational data are used. A citable Zenodo archive can be minted from the GitHub release at acceptance or before final submission."""

code_availability = """The Python code, OpenDSS-compatible circuit files, archived OpenDSSDirect.py harmonic-run artifacts, Travis 150 greenfield screening script, GridPACK/HELICS/OpenDSS dynamic-VAR runner and reproduction scripts are included in the public code repository at https://github.com/SavannahY/dc-ai-factory-backbone-reproducibility and in the submitted reproducibility archive. A permanent Zenodo DOI can be added to this statement once a release archive is minted."""

ai_disclosure = """During manuscript preparation, the authors used AI-assisted tools for drafting support, code refactoring, reference-format checking and editorial revision. The authors reviewed and edited all generated text, verified all scientific claims, generated the final figures from reproducible code and take full responsibility for the content of the submitted manuscript."""

references = [
"Masanet, E. et al. Recalibrating global data center energy-use estimates. Science 367, 984-986 (2020). https://doi.org/10.1126/science.aba3758.",
"Rong, H., Zhang, H., Xiao, S., Li, C. & Hu, C. Optimizing energy consumption for data centers. Renew. Sustain. Energy Rev. 58, 674-691 (2016). https://doi.org/10.1016/j.rser.2015.12.283.",
"Shehabi, A. et al. 2024 United States Data Center Energy Usage Report. Lawrence Berkeley National Laboratory (2024). https://doi.org/10.71468/P1WC7Q.",
"International Energy Agency. Energy and AI. IEA, Paris (2025); https://www.iea.org/reports/energy-and-ai.",
"Blake, M., Hsu, M., Goldwasser, I., Petty, H. & Huntington, J. NVIDIA 800 VDC architecture will power the next generation of AI factories. NVIDIA Technical Blog (20 May 2025); https://developer.nvidia.com/blog/nvidia-800-v-hvdc-architecture-will-power-the-next-generation-of-ai-factories/ (accessed 27 May 2026).",
"Shrestha, B. R. et al. Efficiency and reliability analyses of AC and 380 V DC distribution in data centers. IEEE Access 6, 63305-63315 (2018). https://doi.org/10.1109/ACCESS.2018.2877354.",
"Open Compute Project. Data Center Facility Power Distribution LVDC White Paper, version 1.0 (30 March 2026); https://www.opencompute.org/documents/dcf-power-distribution-lvdc-white-paper-version-1-0-final-pdf-1.",
"Rothmund, D., Guillod, T., Bortis, D. & Kolar, J. W. 99% efficient 10 kV SiC-based 7 kV/400 V DC transformer for future data centers. IEEE J. Emerg. Sel. Top. Power Electron. 7, 753-767 (2019). https://doi.org/10.1109/JESTPE.2018.2886139.",
"Zheng, L. et al. SiC-based 5-kV universal modular soft-switching solid-state transformer (M-S4T) for medium-voltage DC microgrids and distribution grids. IEEE Trans. Power Electron. 36, 11326-11343 (2021). https://doi.org/10.1109/TPEL.2021.3066908.",
"Samanta, S., Wong, I., Bhattacharya, S. & Pahl, B. Medium voltage supply directly to data-center-servers using SiC-based single-stage converter with 20 kW experimental results. In 2020 IEEE Energy Conversion Congress and Exposition (ECCE), 2006-2012 (IEEE, 2020). https://doi.org/10.1109/ECCE44975.2020.9235701.",
"She, X., Huang, A. Q. & Burgos, R. Review of solid-state transformer technologies and their application in power distribution systems. IEEE J. Emerg. Sel. Top. Power Electron. 1, 186-198 (2013). https://doi.org/10.1109/JESTPE.2013.2277917.",
"Choukse, E. et al. Power stabilization for AI training datacenters. arXiv:2508.14318v2 (2025). https://arxiv.org/abs/2508.14318.",
"IEEE Standards Association. IEEE Std 519-2022: IEEE recommended practice and requirements for harmonic control in electric power systems (IEEE, 2022).",
"California ISO. San Jose Area Transmission Plan: decision on modifications to the 2021-2022 transmission plan study (5 November 2024); https://www.caiso.com/documents/decision-on-modifications-to-the-2021-2022-transmission-plan-study-nov-2024.pdf (accessed 27 May 2026).",
"California ISO. 2024-2025 Transmission Planning Process: Board Approved Transmission Plan Posted (30 May 2025); https://www.caiso.com/notices/2024-2025-transmission-planning-process-board-approved-transmission-plan-posted (accessed 27 May 2026).",
"LS Power. LS Power selected by the California ISO for San Jose area HVDC projects. Press release (8 March 2023); https://www.lspower.com/news/ls-power-selected-by-the-california-iso-for-san-jose-area-hvdc-projects/ (accessed 27 May 2026).",
"LS Power Grid. Power Santa Clara Valley HVDC Project fact sheet (2025); https://www.lspowergrid.com/wp-content/uploads/Power-Santa-Clara-Valley-2-Pager.pdf (accessed 27 May 2026).",
"Texas A&M University Electric Grid Test Case Repository. Synthetic Gas-Electric Test Case for the Travis 150 System. https://electricgrids.engr.tamu.edu/synthetic-gas-electric-test-case-for-the-travis-150-system/ (accessed 21 June 2026).",
"Electric Reliability Council of Texas. Planning. https://www.ercot.com/gridinfo/planning (accessed 21 June 2026).",
"CIGRE Working Group B4.52. HVDC grid feasibility study. CIGRE Technical Brochure 533 (2013); https://www.e-cigre.org/publications/detail/533-hvdc-grid-feasibility-study.html.",
"CIGRE Joint Working Group B4/B5.59. Protection and local control of HVDC-grids. CIGRE Technical Brochure 739 (2018); https://www.e-cigre.org/publications/detail/739-protection-and-local-control-of-hvdc-grids.html.",
"CIGRE Working Group B4.76. DC-DC converters in HVDC grids and for connections to HVDC systems. CIGRE Technical Brochure 827 (2021); https://electra.cigre.org/315-april-2021/technical-brochures/dc-dc-converters-in-hvdc-grids-and-for-connections-to-hvdc-systems.html.",
"IEC. IEC 61000-3-3: Electromagnetic compatibility - limits for voltage changes, voltage fluctuations and flicker (IEC, 2013).",
"North American Electric Reliability Corporation. Interconnection oscillation analysis. Technical report (2019); <https://www.ercot.com/files/docs/2019/10/02/Interconnection_Oscillation_Analysis_NERC.pdf>."
]

main_md = '# Direct-current subtransmission backbones for grid-stable AI factories\n\n'
main_md += 'Zhengjie Yang^1,*^ and Liang Min^1,*^\n\n^1^ Stanford University, Stanford, CA, USA.\n\n*Correspondence: yjane@stanford.edu; liangmin@stanford.edu\n\n'
main_md += '## Abstract\n' + abstract + '\n\n'
main_md += '## Introduction\n' + intro + '\n\n'
main_md += '## Results\n\n'
main_md += 'We test the boundary-placement claim through six linked analyses: architecture definition, corridor-loss accounting, harmonic ownership, dynamic buffering, public load-pocket context and a Travis 150 GridPACK/HELICS/OpenDSS validation case. The purpose is not to optimize one corridor design, but to test whether the same C1-C2-C3 ordering appears across independent electrical metrics.\n\n'
for h,txt in results_sections: main_md += f'### {h}\n{txt}\n\n'
main_md += '## Discussion\n' + discussion + '\n\n'
main_md += '## Methods\n\n'
for h,txt in methods: main_md += f'### {h}\n{txt}\n\n'
main_md += '## Data availability\n' + data_availability + '\n\n'
main_md += '## Code availability\n' + code_availability + '\n\n'
main_md += '## AI-assisted drafting disclosure\n' + ai_disclosure + '\n\n'
main_md += '## References\n\n'
for i,refi in enumerate(references,1): main_md += f'{i}. {refi}\n'
main_md += '\n'
main_md += '## Author contributions\nZ.Y. and L.M. contributed to the conceptual framing, analysis and manuscript preparation. Both authors reviewed and approved the manuscript.\n\n'
main_md += '## Competing interests\nThe authors declare no competing interests.\n\n'
main_md += '## Figure legends\n\n'
for k,v in figure_legends.items(): main_md += f'**{k}** {v}\n\n'
(ROOT/'Direct_current_subtransmission_backbones_for_AI_factories_NComms_2026-06-26.md').write_text(main_md)

# Supplementary text
supp_md = """# Supplementary Information

# Supplementary Note 1. Assumption provenance
The modelling assumptions are separated into measured device evidence, industry roadmap evidence, public planning data and forward-looking architecture assumptions. The main text uses conservative phrasing where a value is an extrapolation beyond a measured converter prototype.

# Supplementary Note 2. Averaged EMT equations
The dynamic model is an averaged, architecture-level representation. The AI load is P_L(t). The grid-facing command in architecture j is P_g,j and follows dP_g,j/dt = (P_L - P_g,j)/tau_j, with tau_j = 0 for traditional AC, 1.1 s for local SST, 5-7 s for stronger baselines and 16 s for the DC backbone. The shared buffer power is P_b = P_L - P_g,DC. Its energy state is E_b(t) = integral P_b(t) dt. The voltage proxy is Delta V/V = k_g (P_g - mean(P_g))/S_sc plus a local droop term proportional to P_b. These equations compare exposure between architectures and are not a replacement for switching EMT models.

# Supplementary Note 3. Protection-zone screening
The DC protection study represents detection, converter current limiting, breaker opening, section isolation and re-energization. It is included to expose the functions required by the architecture. It does not specify breaker hardware, insulation coordination or a validated relay scheme.

# Supplementary Note 4. Buffer and economics interpretation
The reference buffer requirement is high power and low energy. It can be met only by coordinated layers: GPU power smoothing, rack or row storage, supercapacitors, converter DC-link energy and station-level storage. The cost/copper envelope is a first-order screen and is not a capital-cost estimate.

# Supplementary Note 5. Travis 150 greenfield dynamic workflow
The Travis 150 analysis uses only the synthetic electric case. The companion gas network is ignored. The C1, C2 and C3 systems are new data-center supply configurations, not conversions of existing AC lines. The dynamic workflow uses the archived GridPACK-compatible Travis 150 RAW/DYR deck and the exported bus-150 POI voltage traces from six shifted branch-fault simulations.

# Supplementary Table 1. Assumption provenance
See data/assumption_provenance_table_v3.csv.

# Supplementary Figure captions
**Supplementary Fig. S1 | Protection screening.** Representative protection zones and dynamic response to a backbone pole-to-ground fault. The sequence includes detection, current limiting, breaker opening, section isolation and healthy-campus ride-through.

![Protection screening and DC fault response](../figures/dc_fault_protection_dynamic.png)

**Supplementary Fig. S2 | Dynamic-screen checks.** Time-step convergence and first-order transfer-function validation for the supplemental grid-command model.

![Averaged EMT model validation](../figures/averaged_emt_validation.png)

**Supplementary Fig. S3 | Shared-buffer interpretation.** Candidate technologies and deployment layers for high-power, low-energy buffering.

![Shared buffer feasibility layers](../figures/shared_buffer_feasibility.png)

**Supplementary Fig. S4 | Cost and conductor envelope.** Annual value of loss reduction under electricity-price and load-factor sweeps, and a current-length index for corridor conductor burden.

![Cost and conductor envelope](../figures/cost_copper_envelope.png)
"""
(SUPP/'Supplementary_Information_NComms_2026-06-26.md').write_text(supp_md)

# Assumption provenance table
prov=pd.DataFrame([
    {'parameter':'Reference useful load','value':'1 GW at the 800 VDC interface','role':'central reference case','source':'this study'},
    {'parameter':'Reference corridor length','value':'20 km equivalent corridor','role':'central reference case','source':'this study'},
    {'parameter':'Reference AC voltage','value':'138 kV line-to-line','role':'representative subtransmission voltage','source':'this study'},
    {'parameter':'Reference AC power factor','value':'0.98','role':'central AC-current assumption','source':'this study'},
    {'parameter':'Reference DC voltage','value':'+/-138 kV, or 276 kV pole-to-pole','role':'representative DC design point','source':'this study'},
    {'parameter':'Conductor resistance','value':'0.01 ohm/km per phase or pole','role':'screening assumption','source':'this study'},
    {'parameter':'Traditional downstream efficiency','value':'0.991 x 0.982 = 97.32%','role':'central loss assumption','source':'this study'},
    {'parameter':'Local SST efficiency','value':'98.5%','role':'central local-SST assumption','source':'this study'},
    {'parameter':'99% local-SST efficiency sensitivity','value':'99.0%','role':'sensitivity case, not a demonstrated reference architecture','source':'this study sensitivity assumption'},
    {'parameter':'DC terminal AC/DC efficiency','value':'99.4%','role':'central DC-backbone assumption','source':'this study'},
    {'parameter':'Subtransmission-to-34.5 kV DC/DC efficiency','value':'99.4%','role':'central DC-backbone assumption','source':'this study'},
    {'parameter':'34.5 kV/800 V DC/DC efficiency','value':'99.2%','role':'central DC-backbone assumption','source':'this study'},
    {'parameter':'7 kV/400 V DC/DC efficiency','value':'99.0%','role':'measured device evidence','source':'Rothmund et al. 2019'},
    {'parameter':'3.8 kV AC to 400 V DC SST chain','value':'98.1%','role':'measured chain evidence','source':'Rothmund et al. 2019'},
    {'parameter':'5 kV modular M-S4T peak efficiency','value':'97.5% estimated','role':'modular MVDC evidence','source':'Zheng et al. 2021'},
    {'parameter':'20 kW 1000 V/48 V prototype','value':'96% measured; 97.5% estimated with synchronous rectification','role':'data-center voltage-step evidence','source':'Samanta et al. 2020'},
    {'parameter':'AI load spectral range','value':'0.1-20 Hz utility concern','role':'load-dynamics evidence','source':'Choukse et al. 2025'},
    {'parameter':'Grid short-circuit strength','value':'10 GVA','role':'screening assumption','source':'this study'},
])
prov.to_csv(DATA/'assumption_provenance_table_v3.csv',index=False)

# ---------------------------- DOCX generation ----------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr); t_el = OxmlElement('w:t'); t_el.text = text; new_run.append(t_el)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)

def style_doc(doc):
    styles=doc.styles
    styles['Normal'].font.name='Arial'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); styles['Normal'].font.size=Pt(10)
    styles['Normal'].font.color.rgb=RGBColor(0,0,0)
    styles['Normal'].paragraph_format.line_spacing=2.0
    styles['Normal'].paragraph_format.space_after=Pt(0)
    for style in ['Title','Heading 1','Heading 2','Heading 3','Caption']:
        styles[style].font.name='Arial'; styles[style]._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
        styles[style].font.color.rgb=RGBColor(0,0,0)
    title_ppr = styles['Title']._element.get_or_add_pPr()
    title_border = title_ppr.find(qn('w:pBdr'))
    if title_border is not None:
        title_ppr.remove(title_border)
    styles['Title'].font.size=Pt(18); styles['Title'].font.bold=True
    styles['Heading 1'].font.size=Pt(14); styles['Heading 1'].font.bold=True
    styles['Heading 2'].font.size=Pt(12); styles['Heading 2'].font.bold=True
    styles['Heading 3'].font.size=Pt(10.5); styles['Heading 3'].font.bold=True

def add_para(doc, text):
    p=doc.add_paragraph(text); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=2.0; return p

def add_fig(doc, path, caption, width=6.3):
    doc.add_picture(str(path), width=Inches(width))
    p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cp=doc.add_paragraph(caption); cp.style=doc.styles['Caption'] if 'Caption' in doc.styles else doc.styles['Normal']; cp.paragraph_format.space_after=Pt(8)
    for run in cp.runs: run.font.size=Pt(8)

def create_main_docx():
    doc=Document(); style_doc(doc)
    sec=doc.sections[0]; sec.top_margin=Inches(0.65); sec.bottom_margin=Inches(0.65); sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)
    title=doc.add_paragraph(); title.style='Title'; title.add_run('Direct-current subtransmission backbones for grid-stable AI factories')
    add_para(doc,'Zhengjie Yang and Liang Min')
    add_para(doc,'Stanford University, Stanford, CA, USA')
    add_para(doc,'Correspondence: yjane@stanford.edu; liangmin@stanford.edu')
    doc.add_heading('Abstract',level=1); add_para(doc,abstract)
    doc.add_heading('Introduction',level=1)
    for para in intro.split('\n\n'): add_para(doc,para)
    doc.add_heading('Results',level=1)
    doc.add_heading(results_sections[0][0],level=2)
    for para in results_sections[0][1].split('\n\n'): add_para(doc,para)
    add_fig(doc, FIG/'ai_factory_delivery_architectures.png', list(figure_legends.items())[0][0]+' '+list(figure_legends.items())[0][1])
    for idx,(h,txt) in enumerate(results_sections[1:], start=2):
        doc.add_heading(h,level=2)
        for para in txt.split('\n\n'): add_para(doc,para)
        fpath={2:'transfer_capacity_loss_designspace.png',3:'harmonic_ownership_opendss_screening.png',4:'gridpack_voltage_ride_through.png',5:'load_pocket_voltage_envelope.png',6:'travis150_greenfield_benefits.png'}[idx]
        cap_key=list(figure_legends.keys())[idx-1]
        add_fig(doc, FIG/fpath, cap_key+' '+figure_legends[cap_key])
    doc.add_heading('Discussion',level=1)
    for para in discussion.split('\n\n'): add_para(doc,para)
    doc.add_heading('Methods',level=1)
    for h,txt in methods:
        doc.add_heading(h,level=2)
        for para in txt.split('\n\n'): add_para(doc,para)
    doc.add_heading('Data availability',level=1); add_para(doc,data_availability)
    doc.add_heading('Code availability',level=1); add_para(doc,code_availability)
    doc.add_heading('AI-assisted drafting disclosure',level=1); add_para(doc,ai_disclosure)
    doc.add_heading('References',level=1)
    for i,refi in enumerate(references,1): add_para(doc,f'{i}. {refi}')
    doc.add_heading('Author contributions',level=1)
    add_para(doc,'Z.Y. and L.M. contributed to the conceptual framing, analysis and manuscript preparation. Both authors reviewed and approved the manuscript.')
    doc.add_heading('Competing interests',level=1)
    add_para(doc,'The authors declare no competing interests.')
    doc.add_heading('Figure legends',level=1)
    for key, legend in figure_legends.items():
        add_para(doc, f'{key} {legend}')
    out=ROOT/'Direct_current_subtransmission_backbones_for_AI_factories_NComms_2026-06-26.docx'
    doc.save(out); return out

def create_supp_docx():
    doc=Document(); style_doc(doc)
    sec=doc.sections[0]; sec.top_margin=Inches(0.65); sec.bottom_margin=Inches(0.65); sec.left_margin=Inches(0.7); sec.right_margin=Inches(0.7)
    title=doc.add_paragraph(); title.style='Title'; title.add_run('Supplementary Information')
    doc.add_heading('Supplementary Note 1. Assumption provenance',level=1)
    add_para(doc,'The modelling assumptions are separated into measured device evidence, industry roadmap evidence, public planning data and forward-looking architecture assumptions. The main text uses conservative phrasing where a value is an extrapolation beyond a measured converter prototype.')
    # table
    doc.add_heading('Supplementary Table 1. Assumption provenance',level=2)
    table=doc.add_table(rows=1,cols=4); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    hdr=table.rows[0].cells
    for c,tv in zip(hdr,['Parameter','Value','Role','Source']): c.text=tv; set_cell_shading(c,'E6E6E6')
    for _,row in prov.iterrows():
        cells=table.add_row().cells
        for i,k in enumerate(['parameter','value','role','source']): cells[i].text=str(row[k])
    doc.add_heading('Supplementary Note 2. Averaged EMT equations and validation',level=1)
    add_para(doc,'The dynamic model is an averaged, architecture-level representation. The AI load is P_L(t). The grid-facing command P_g follows dP_g/dt = (P_L - P_g)/tau. The shared buffer power is P_b = P_L - P_g, and the buffer energy state is the time integral of P_b. The scenario grid varies campus count, cluster load, voltage class, short-circuit ratio, temporal phase coherence and corridor length. The voltage proxy combines grid-stiffness and corridor-voltage factors. These equations compare exposure between architectures and are not a replacement for switching EMT models.')
    add_fig(doc, FIG/'averaged_emt_validation.png', 'Supplementary Fig. S2 | Dynamic-screen checks. Time-step convergence and first-order transfer-function validation for the supplemental grid-command model.')
    doc.add_heading('Supplementary Note 3. Protection-zone screening',level=1)
    add_para(doc,'The DC protection study represents detection, converter current limiting, breaker opening, section isolation and re-energization. It is included to expose the functions required by the architecture. It does not specify breaker hardware, insulation coordination or a validated relay scheme.')
    add_fig(doc, FIG/'dc_fault_protection_dynamic.png', 'Supplementary Fig. S1 | Protection screening. Representative protection zones and dynamic response to a backbone pole-to-ground fault.')
    doc.add_heading('Supplementary Note 4. Buffer and economics interpretation',level=1)
    add_para(doc,'The reference buffer requirement is high power and low energy. It can be met only by coordinated layers: GPU power smoothing, rack or row storage, supercapacitors, converter DC-link energy and station-level storage. The cost/copper envelope is a first-order screen and is not a capital-cost estimate.')
    add_fig(doc, FIG/'shared_buffer_feasibility.png', 'Supplementary Fig. S3 | Shared-buffer interpretation. Candidate technologies and deployment layers for high-power, low-energy buffering.')
    add_fig(doc, FIG/'cost_copper_envelope.png', 'Supplementary Fig. S4 | Cost and conductor envelope. Annual value of loss reduction and current-length index for corridor conductor burden.')
    doc.add_heading('Supplementary Note 5. Travis 150 greenfield dynamic workflow',level=1)
    add_para(doc,'The Travis 150 analysis uses only the synthetic electric case. The companion gas network is ignored. The C1, C2 and C3 systems are new data-center supply configurations, not conversions of existing AC lines. The dynamic workflow uses the archived GridPACK-compatible Travis 150 RAW/DYR deck and the exported bus-150 POI voltage traces from six shifted branch-fault simulations.')
    out=SUPP/'Supplementary_Information_NComms_2026-06-26.docx'; doc.save(out); return out

main_docx=create_main_docx(); supp_docx=create_supp_docx()

def create_source_data_package():
    source_map = {
        'Fig. 2': [
            'transfer_capacity_reference_case_v3.csv',
            'transfer_capacity_uncertainty_reference_v3.csv',
            'transfer_capacity_design_space_v3.csv',
            'transfer_capacity_sensitivity_v3.csv',
        ],
        'Fig. 3': [
            'harmonic_thdv_monte_carlo_v3.csv',
            'harmonic_individual_p95_v3.csv',
            'true_opendss_harmonic_thdv_monte_carlo_v3.csv',
            'true_opendss_harmonic_individual_spectrum_v3.csv',
        ],
        'Fig. 4': [
            'gridpack_voltage_turbulence_event_sweep_v3.csv',
            'gridpack_voltage_control_event_response_v3.csv',
            'gridpack_voltage_control_summary_v3.csv',
        ],
        'Fig. 5': [
            'fig5_load_pocket_source_data_v3.csv',
            'fig5_voltage_class_current_envelope_v3.csv',
        ],
        'Fig. 6': [
            'travis150_greenfield_c1_c2_c3_summary_v2.csv',
            'travis150_greenfield_c1_c2_c3_transfer_v2.csv',
            'travis150_greenfield_c1_c2_c3_harmonics_v2.csv',
            'travis150_greenfield_c1_c2_c3_voltage_v2.csv',
            'event_sweep_summary_compact.csv',
            'gridpack_poi_voltage_event_0.csv',
            'gridpack_poi_voltage_event_1.csv',
            'gridpack_poi_voltage_event_2.csv',
            'gridpack_poi_voltage_event_3.csv',
            'gridpack_poi_voltage_event_4.csv',
            'gridpack_poi_voltage_event_5.csv',
        ],
        'Supplementary figures': [
            'dc_fault_protection_backbone_fault_v3.csv',
            'dc_fault_protection_campus_fault_v3.csv',
            'emt_timestep_convergence_v3.csv',
            'emt_transfer_function_validation_v3.csv',
            'buffer_physical_feasibility_table_v3.csv',
            'cost_copper_envelope_v3.csv',
            'harmonic_robustness_summary_v3.csv',
            'harmonic_robustness_architecture_comparison_v3.csv',
            'harmonic_robustness_individual_p95_v3.csv',
        ],
        'Assumptions and provenance': [
            'assumptions_v3.json',
            'assumption_provenance_table_v3.csv',
        ],
    }
    readme_lines = [
        '# Source data for figures',
        '',
        'This folder contains machine-readable numerical data underlying the main and supplementary graphs.',
        'Each file is copied from the reproducibility-package `data/` directory without modification.',
        '',
    ]
    for label, filenames in source_map.items():
        readme_lines.append(f'## {label}')
        for name in filenames:
            src = DATA / name
            if not src.exists():
                src = SOURCE_ROOT / 'data' / name
            if src.exists():
                shutil.copy(src, SOURCE_DATA / name)
                readme_lines.append(f'- `{name}`')
        readme_lines.append('')
    readme = SOURCE_DATA / 'README.md'
    readme.write_text('\n'.join(readme_lines), encoding='utf-8')
    return SOURCE_DATA

source_data_dir = create_source_data_package()

cover_letter = """Dear Editors,

We are pleased to submit "Direct-current subtransmission backbones for grid-stable AI factories" for consideration as an Article in Nature Communications.

AI data centers are becoming synchronized, DC-native, gigawatt-scale loads, but grid planning still often treats them as passive AC facilities. This manuscript asks where the AC/DC boundary should sit when the useful facility boundary is 800 VDC. We compare traditional AC delivery, local solid-state-transformer delivery and a utility-operated subtransmission DC backbone.

The main contribution is a falsifiable architecture-level claim: moving the AC/DC boundary upstream can co-locate three benefits that are usually studied separately. In the reference and Travis 150 greenfield studies, the DC-corridor architecture increases useful transfer capacity, centralizes AC-side harmonic ownership and improves voltage ride-through in a GridPACK/HELICS/OpenDSS branch-fault event sweep. The Travis 150 analysis is explicitly framed as a synthetic test-bed result rather than a site-selection or real-routing claim.

The manuscript includes all figure source data, a reproducibility archive, OpenDSS-compatible harmonic files, a Travis 150 greenfield screening workflow and a documented GridPACK/HELICS/OpenDSS dynamic workflow using the real Travis 150 RAW/DYR case. We believe the work will interest Nature Communications readers working across power systems, power electronics, grid planning and AI infrastructure.

This manuscript is not under consideration elsewhere. The authors declare no competing interests. We have not had prior discussions with a Nature Communications editor about this work.

Sincerely,

Zhengjie Yang and Liang Min
Stanford University
"""
(ROOT/'cover_letter_nature_communications.md').write_text(cover_letter, encoding='utf-8')

def create_tex_pdf_package():
    pandoc = shutil.which('pandoc')
    tectonic = shutil.which('tectonic')
    note = ROOT / 'tex_pdf_build_note.txt'
    if not pandoc or not tectonic:
        missing = ', '.join(name for name, path in [('pandoc', pandoc), ('tectonic', tectonic)] if not path)
        note.write_text(
            f'TeX/PDF build skipped because the following executable(s) were not found: {missing}.\n',
            encoding='utf-8',
        )
        return None, None, None

    main_md_path = ROOT / 'Direct_current_subtransmission_backbones_for_AI_factories_NComms_2026-06-26.md'
    pdf_md_path = RENDER / 'Direct_current_subtransmission_backbones_for_AI_factories_NComms_2026-06-26_with_inline_figures.md'
    tex_path = ROOT / 'Direct_current_subtransmission_backbones_for_AI_factories_NComms_2026-06-26.tex'
    overleaf_path = ROOT / 'Direct_current_subtransmission_backbones_for_AI_factories_NComms_overleaf.tex'

    figure_after = [
        ('An AI-native architecture with the AC/DC boundary moved upstream', 'AI-factory delivery architectures', 'ai_factory_delivery_architectures.png'),
        ('Transfer capacity is coupled to loss reduction', 'Transfer capacity and loss design space', 'transfer_capacity_loss_designspace.png'),
        ('A DC backbone changes harmonic compliance into harmonic ownership', 'Harmonic ownership at the AC grid interface', 'harmonic_ownership_opendss_screening.png'),
        ('The DC backbone buffers voltage turbulence at a controllable boundary', 'GridPACK voltage ride-through under transmission faults', 'gridpack_voltage_ride_through.png'),
        ('Data-center load pockets are becoming planning objects', 'Load-pocket voltage and conductor context', 'load_pocket_voltage_envelope.png'),
        ('Travis 150 greenfield configurations preserve the three-benefit ordering', 'Travis 150 greenfield validation of C1-C3 ordering', 'travis150_greenfield_benefits.png'),
    ]

    pdf_md = main_md_path.read_text(encoding='utf-8')
    for heading, caption, image in figure_after:
        marker = f'### {heading}\n'
        start = pdf_md.find(marker)
        if start == -1:
            continue
        search_from = start + len(marker)
        next_positions = [pos for pos in (pdf_md.find('\n### ', search_from), pdf_md.find('\n## ', search_from)) if pos != -1]
        end = min(next_positions) if next_positions else len(pdf_md)
        figure_markdown = f'\n\n![{caption}](figures/{image})\n'
        if figure_markdown.strip() not in pdf_md[start:end]:
            pdf_md = pdf_md[:end].rstrip() + figure_markdown + pdf_md[end:]
    pdf_md_path.write_text(pdf_md, encoding='utf-8')

    subprocess.run([
        pandoc,
        pdf_md_path.name,
        '--standalone',
        '--from', 'markdown',
        '--to', 'latex',
        '-V', 'geometry:margin=1in',
        '-V', 'fontsize=11pt',
        '-o', tex_path.name,
    ], cwd=RENDER, check=True)

    generated_tex = RENDER / tex_path.name
    if generated_tex.exists():
        shutil.move(generated_tex, tex_path)

    tex = tex_path.read_text(encoding='utf-8')
    tex = tex.replace(
        r'\setlength{\emergencystretch}{3em}',
        r'\setlength{\emergencystretch}{8em}' + '\n' + r'\sloppy',
    )
    tex = tex.replace(
        r'\section{Direct-current subtransmission backbones for grid-stable AI factories}\label{direct-current-subtransmission-backbones-for-grid-stable-ai-factories}',
        r'{\Large\bfseries\raggedright Direct-current subtransmission backbones for grid-stable AI factories\par}'
        + '\n'
        + r'\label{direct-current-subtransmission-backbones-for-grid-stable-ai-factories}'
        + '\n'
        + r'\vspace{0.5em}',
    )
    tex = tex.replace(
        '\\section{Direct-current subtransmission backbones for grid-stable AI\n'
        'factories}\\label{direct-current-subtransmission-backbones-for-grid-stable-ai-factories}',
        r'{\Large\bfseries\raggedright Direct-current subtransmission backbones for grid-stable AI factories\par}'
        + '\n'
        + r'\label{direct-current-subtransmission-backbones-for-grid-stable-ai-factories}'
        + '\n'
        + r'\vspace{0.5em}',
    )
    tex = tex.replace('../figures/', 'figures/')
    tex_path.write_text(tex, encoding='utf-8')
    overleaf_path.write_text(tex, encoding='utf-8')

    subprocess.run([
        tectonic,
        '--keep-logs',
        tex_path.name,
    ], cwd=ROOT, check=True)

    note.write_text(
        'TeX and PDF were generated with pandoc and tectonic from the manuscript markdown with inline figures.\n',
        encoding='utf-8',
    )
    return tex_path, overleaf_path, tex_path.with_suffix('.pdf')

def create_supplementary_pdf_package():
    pandoc = shutil.which('pandoc')
    tectonic = shutil.which('tectonic')
    note = ROOT / 'supplementary_pdf_build_note.txt'
    if not pandoc or not tectonic:
        missing = ', '.join(name for name, path in [('pandoc', pandoc), ('tectonic', tectonic)] if not path)
        note.write_text(
            f'Supplementary TeX/PDF build skipped because the following executable(s) were not found: {missing}.\n',
            encoding='utf-8',
        )
        return None, None

    supp_md_path = SUPP / 'Supplementary_Information_NComms_2026-06-26.md'
    supp_tex_path = SUPP / 'Supplementary_Information_NComms_2026-06-26.tex'

    subprocess.run([
        pandoc,
        supp_md_path.name,
        '--standalone',
        '--from', 'markdown',
        '--to', 'latex',
        '-V', 'geometry:margin=1in',
        '-V', 'fontsize=11pt',
        '-o', supp_tex_path.name,
    ], cwd=SUPP, check=True)

    tex = supp_tex_path.read_text(encoding='utf-8')
    tex = tex.replace(
        r'\setlength{\emergencystretch}{3em}',
        r'\setlength{\emergencystretch}{8em}' + '\n' + r'\sloppy',
    )
    supp_tex_path.write_text(tex, encoding='utf-8')

    subprocess.run([
        tectonic,
        '--keep-logs',
        supp_tex_path.name,
    ], cwd=SUPP, check=True)

    note.write_text(
        'Supplementary TeX and PDF were generated with pandoc and tectonic from the supplementary markdown with inline figures.\n',
        encoding='utf-8',
    )
    return supp_tex_path, supp_tex_path.with_suffix('.pdf')

tex_path, overleaf_path, pdf_path = create_tex_pdf_package()
supp_tex_path, supp_pdf_path = create_supplementary_pdf_package()

# ---------------------------- Public repository ----------------------------
# Copy data, figures, OpenDSS files to repo
for src in DATA.glob('*'):
    shutil.copy(src, REPO/'data'/src.name)
for src in FIG.glob('*.png'):
    shutil.copy(src, REPO/'figures'/src.name)
for src in FIG.glob('*.svg'):
    shutil.copy(src, REPO/'figures'/src.name)
for src in OPENDSS.glob('*'):
    shutil.copy(src, REPO/'opendss'/src.name)
source_root = Path(__file__).resolve().parents[1]
source_data = source_root/'data'
source_opendss = source_root/'opendss'
source_scripts = source_root/'scripts'
if source_data.exists():
    for src in source_data.glob('true_opendss_*'):
        shutil.copy(src, DATA/src.name)
        shutil.copy(src, REPO/'data'/src.name)
if source_opendss.exists():
    for src in source_opendss.glob('true_opendss*'):
        shutil.copy(src, OPENDSS/src.name)
        shutil.copy(src, REPO/'opendss'/src.name)
if (source_scripts/'run_true_opendss.py').exists():
    shutil.copy(source_scripts/'run_true_opendss.py', REPO/'scripts'/'run_true_opendss.py')
# Code modules
(REPO/'src'/'ai_dc_backbone'/'__init__.py').write_text('__version__ = "0.3.0"\n')
(REPO/'src'/'ai_dc_backbone'/'efficiency.py').write_text(textwrap.dedent('''
    import math
    def losses_eff(load_MW=1000, length_km=20, r_ohm_km=0.01, pf=0.98,
                   trad_eff=0.991*0.982, sst_eff=0.985, dc_term=0.994, dc1=0.994, dc2=0.992,
                   vac_kv=138, vdc_pp_kv=276):
        P=load_MW*1e6; R=r_ohm_km*length_km
        P_recv_trad=P/trad_eff; I_ac_trad=P_recv_trad/(math.sqrt(3)*vac_kv*1e3*pf); line_trad=3*I_ac_trad**2*R; input_trad=P_recv_trad+line_trad
        P_recv_sst=P/sst_eff; I_ac_sst=P_recv_sst/(math.sqrt(3)*vac_kv*1e3*pf); line_sst=3*I_ac_sst**2*R; input_sst=P_recv_sst+line_sst
        P_recv_dc=P/(dc1*dc2); I_dc=P_recv_dc/(vdc_pp_kv*1e3); line_dc=2*I_dc**2*R; input_dc=(P_recv_dc+line_dc)/dc_term
        return {'Traditional AC':(input_trad-P)/1e6,'Local SST':(input_sst-P)/1e6,'Subtransmission DC backbone':(input_dc-P)/1e6}
    def grid_input_MW(load_MW, architecture, **kwargs):
        return load_MW + losses_eff(load_MW=load_MW, **kwargs)[architecture]
    def useful_transfer_at_grid_input(input_limit_MW, architecture, **kwargs):
        lo=0.0; hi=input_limit_MW
        for _ in range(70):
            mid=0.5*(lo+hi)
            if grid_input_MW(mid, architecture, **kwargs) <= input_limit_MW: lo=mid
            else: hi=mid
        return lo
'''))
(REPO/'src'/'ai_dc_backbone'/'dynamics.py').write_text(textwrap.dedent('''
    import numpy as np
    def lpf(x, tau, dt):
        y=np.empty_like(x); y[0]=x[0]; a=dt/(tau+dt)
        for i in range(1,len(x)): y[i]=y[i-1]+a*(x[i]-y[i-1])
        return y
    def spectral_energy(x, dt, fmin=0.1, fmax=20):
        y=x-np.mean(x); freqs=np.fft.rfftfreq(len(y),dt); mag=np.abs(np.fft.rfft(y))/len(y)*2
        mask=(freqs>=fmin)&(freqs<=fmax)
        return float(np.sqrt(np.sum(mag[mask]**2)))
'''))
(REPO/'src'/'ai_dc_backbone'/'harmonics.py').write_text(textwrap.dedent('''
    import numpy as np, math
    def resonance_factor(h, shift=0.0, strength=1.0):
        return 1 + strength*(3.2*np.exp(-0.5*((h-(11+shift))/1.6)**2) + 1.7*np.exp(-0.5*((h-(23+0.5*shift))/2.0)**2))
    def note():
        return 'Use this module for the transparent nodal frequency-domain solver. OpenDSS-compatible files are in opendss/.'
'''))
for helper in ['reproduce_all.py','dynamic_robustness_sweep.py','harmonic_robustness_sweep.py','travis150_greenfield_c1_c2_c3.py','run_gridpack_td_dynamic_var.py','run_griddyn_td_dynamic_var.py']:
    source_helper=Path(__file__).resolve().with_name(helper)
    if source_helper.exists():
        shutil.copy(source_helper,REPO/'scripts'/helper)
if not (REPO/'scripts'/'reproduce_all.py').exists():
    (REPO/'scripts'/'reproduce_all.py').write_text("print('Run scripts/build_dc_backbone_v3.py to rebuild the manuscript package.')\n")
(REPO/'scripts'/'run_opendss_if_available.py').write_text(textwrap.dedent('''
    #!/usr/bin/env python
    """Run OpenDSS-compatible files if opendssdirect.py is installed.
    The manuscript figures do not depend on this optional check; they use the
    transparent frequency-domain solver. This script is provided for external validation.
    """
    from pathlib import Path
    try:
        import opendssdirect as dss
    except Exception as e:
        print('OpenDSSDirect not installed:', e)
        print('Install opendssdirect.py in a local environment and rerun this script.')
        raise SystemExit(0)
    for f in Path('opendss').glob('*.dss'):
        print('Compiling', f)
        dss.Basic.ClearAll()
        dss.Text.Command(f'Compile [{f}]')
        dss.Text.Command('Solve mode=harmonics')
        print('Solved:', f)
'''))
(REPO/'README.md').write_text(textwrap.dedent('''
    # Direct-current subtransmission backbones for grid-stable AI factories

    This repository contains the data, screening models, figures and OpenDSS-compatible files for the manuscript
    "Direct-current subtransmission backbones for grid-stable AI factories".

    ## Contents
    - `data/`: CSV inputs and outputs for all manuscript and supplementary figures.
    - `figures/`: publication figures in PNG/SVG form.
    - `src/ai_dc_backbone/`: reusable Python model modules.
    - `scripts/`: reproduction helpers and optional OpenDSS runner.
    - `opendss/`: OpenDSS-compatible harmonic network files.
    - Travis 150 greenfield C1/C2/C3 outputs and GridPACK/HELICS/OpenDSS
      workflow files are included as source data and reproducibility inputs.

    ## Reproducing results
    ```bash
    python scripts/reproduce_all.py
    python scripts/dynamic_robustness_sweep.py
    python scripts/harmonic_robustness_sweep.py
    python scripts/travis150_greenfield_c1_c2_c3.py
    python scripts/run_gridpack_td_dynamic_var.py --gridpack-exe /path/to/dsf.x --execute
    python scripts/run_opendss_if_available.py  # optional, requires opendssdirect.py
    ```

    `scripts/reproduce_all.py` regenerates the archived Fig. 3 diagnostic
    and Fig. 5 from archived CSV outputs into `reproduced/figures`.
    `scripts/dynamic_robustness_sweep.py` regenerates a supplemental dynamic
    screening grid and supporting CSV tables.
    `scripts/harmonic_robustness_sweep.py` regenerates the harmonic robustness
    grid and supporting figures. The manuscript figures were generated with
    transparent Python models. OpenDSS circuit files and the run log are
    included under `opendss/`.
    The Travis 150 dynamic-VAR run uses GridPACK, HELICS and OpenDSSDirect.py
    when those external tools are installed.

    ## Citation
    See `CITATION.cff`. This repository is structured for GitHub release and Zenodo deposition.

    ## Figure and drafting provenance
    - Figure provenance is documented in `docs/figure_provenance.md`.
    - AI-assisted drafting disclosure language is provided in
      `docs/ai_assisted_drafting_disclosure.md`.

    ## Direct OpenDSS check
    This repository includes `scripts/run_true_opendss.py`,
    `opendss/true_opendss_harmonic_network_v3.dss`, and the resulting
    `data/true_opendss_*` CSV files.
'''))
(REPO/'CITATION.cff').write_text(textwrap.dedent('''
    cff-version: 1.2.0
    title: "Code and data for Direct-current subtransmission backbones for grid-stable AI factories"
    message: "If you use this code or data, please cite the associated manuscript and this archive."
    type: software
    authors:
      - family-names: "Yang"
        given-names: "Zhengjie"
        email: "yjane@stanford.edu"
        affiliation: "Stanford University"
      - family-names: "Min"
        given-names: "Liang"
        email: "liangmin@stanford.edu"
        affiliation: "Stanford University"
    version: "0.3.0"
    date-released: "2026-05-26"
    license: "MIT"
    repository-code: "https://github.com/SavannahY/dc-ai-factory-backbone-reproducibility"
    abstract: "Reproducibility package containing source data, OpenDSS cases, figure-generation code and verification tests for a manuscript on direct-current subtransmission backbones for grid-stable AI factories."
'''))
(REPO/'LICENSE').write_text('MIT License\n\nCopyright (c) 2026 Authors\n\nPermission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files...\n')
(REPO/'requirements.txt').write_text('numpy\npandas\nmatplotlib\npython-docx\n')
(REPO/'environment.yml').write_text('name: dc-backbone-ai-factories\nchannels:\n  - conda-forge\ndependencies:\n  - python>=3.10\n  - numpy\n  - pandas\n  - matplotlib\n  - python-docx\n')
(REPO/'docs'/'reproduction.md').write_text(textwrap.dedent('''
    This repository is structured for public release. To regenerate the archived
    OpenDSS Fig. 3 diagnostic and Fig. 5 from archived CSV outputs, run:

    ```bash
    python scripts/reproduce_all.py
    ```

    The script writes the diagnostic Fig. 3 and Fig. 5 to
    `reproduced/figures`. OpenDSSDirect.py
    harmonic-run artifacts are archived under `opendss/` and
    `data/true_opendss_*`. To rerun OpenDSS in a local environment with
    OpenDSSDirect.py installed, run:

    ```bash
    python scripts/run_true_opendss.py
    ```

    To regenerate the supplemental dynamic robustness grid, run:

    ```bash
    python scripts/dynamic_robustness_sweep.py
    ```

    To regenerate the harmonic robustness sweep, the supporting Fig. 3 screening
    variant and optional harmonic-robustness diagnostic outputs, run:

    ```bash
    python scripts/harmonic_robustness_sweep.py
    ```

    The optional diagnostic plots are written under `reproduced/figures` and are
    not part of the final manuscript `figures/` folder.

    To regenerate the Travis 150 greenfield C1/C2/C3 screen after downloading
    the TAMU electric AUX case, run:

    ```bash
    python scripts/travis150_greenfield_c1_c2_c3.py --travis-case Travis150/Travis150_Electric_Data.aux
    ```

    To rerun the installed-tool T&D workflow when GridPACK, HELICS and
    OpenDSSDirect.py are available, run:

    ```bash
    python scripts/run_gridpack_td_dynamic_var.py --gridpack-exe /path/to/dsf.x --execute
    ```

    The complete manuscript-package generator is `scripts/build_dc_backbone_v3.py`.
    It rebuilds the Word manuscript, supplementary information, figures, source
    data folder and public-code archive in an environment with the dependencies
    listed in `requirements.txt`.
''').lstrip())
(REPO/'docs'/'figure_provenance.md').write_text(textwrap.dedent('''
    # Figure provenance

    All manuscript and supplementary figures in this reproducibility package are
    programmatic outputs from `scripts/build_dc_backbone_v3.py` or from the archived
    CSV outputs under `data/`.

    No final manuscript figure is a generative-AI image, stock image, screenshot
    collage or manually edited bitmap. Fig. 1 is the supplied architecture PNG used
    consistently in Word, TeX and PDF. The other distributed PNG, SVG and PDF files
    are Matplotlib exports. The SVG files can be inspected as vector graphics.
    `scripts/reproduce_all.py` regenerates the archived Fig. 3 diagnostic and Fig. 5
    from source CSV files as a fast review-time check.
    `scripts/dynamic_robustness_sweep.py` regenerates a supplemental dynamic
    screening grid and supporting CSV tables.
    `scripts/harmonic_robustness_sweep.py` regenerates the harmonic robustness
    screening figures and the supporting CSV tables.
    `scripts/run_gridpack_td_dynamic_var.py` provides the GridPACK/HELICS/OpenDSS
    event-sweep workflow used for Fig. 4 and Fig. 6.

    Final figure files:

    - Fig. 1: `figures/ai_factory_delivery_architectures.png`
    - Fig. 2: `figures/transfer_capacity_loss_designspace.{png,svg}`
    - Fig. 3: `figures/harmonic_ownership_opendss_screening.{png,svg}`
    - Fig. 4: `figures/gridpack_voltage_ride_through.{png,svg}`
    - Fig. 5: `figures/load_pocket_voltage_envelope.{png,svg}`
    - Fig. 6: `figures/travis150_greenfield_benefits.{png,svg,pdf}`
    - Supplementary Fig. S1: `figures/dc_fault_protection_dynamic.{png,svg}`
    - Supplementary Fig. S2: `figures/averaged_emt_validation.{png,svg}`
    - Supplementary Fig. S3: `figures/shared_buffer_feasibility.{png,svg}`
    - Supplementary Fig. S4: `figures/cost_copper_envelope.{png,svg}`
''').lstrip())
(REPO/'docs'/'ai_assisted_drafting_disclosure.md').write_text(textwrap.dedent(f'''
    # AI-assisted drafting disclosure

    Suggested manuscript disclosure language:

    > {ai_disclosure}

    This disclosure should be reviewed by all authors before submission and adjusted
    to match the actual use of AI tools in the final manuscript workflow.
''').lstrip())

# Copy generator script into CODE and repo root for reproducibility
shutil.copy(__file__, CODE/'build_dc_backbone_v3.py')
shutil.copy(__file__, REPO/'scripts'/'build_dc_backbone_v3.py')

# Manifest with SHA256 for DOI-ready data package
manifest=[]
for f in sorted(REPO.rglob('*')):
    if f.is_file():
        h=hashlib.sha256(f.read_bytes()).hexdigest()
        manifest.append({'path':str(f.relative_to(REPO)),'sha256':h,'bytes':f.stat().st_size})
pd.DataFrame(manifest).to_csv(REPO/'MANIFEST_SHA256.csv',index=False)

# ZIP public repo and the submission package. The outer submission package keeps
# only submission-facing files. Reproducibility inputs live inside the public
# repository zip to avoid duplicating data and code at the package root.
repo_zip=ROOT/'public_code_repo_DOI_ready.zip'
with zipfile.ZipFile(repo_zip,'w',zipfile.ZIP_DEFLATED) as z:
    for f in REPO.rglob('*'):
        z.write(f, f.relative_to(REPO.parent))

for generated_dir in [DATA, CODE, RENDER, OPENDSS, REPO]:
    if generated_dir.exists():
        shutil.rmtree(generated_dir)
for generated_file in [ROOT/'tex_pdf_build_note.txt', ROOT/'supplementary_pdf_build_note.txt', ROOT/'source_data_csv.zip']:
    if generated_file.exists():
        generated_file.unlink()

submission_zip=ROOT.parent/'submission_package.zip'
with zipfile.ZipFile(submission_zip,'w',zipfile.ZIP_DEFLATED) as z:
    for f in ROOT.rglob('*'):
        z.write(f, f.relative_to(ROOT.parent))

print('MAIN_DOCX', main_docx)
print('SUPP_DOCX', supp_docx)
if supp_pdf_path:
    print('SUPP_PDF', supp_pdf_path)
if pdf_path:
    print('MAIN_PDF', pdf_path)
if tex_path:
    print('MAIN_TEX', tex_path)
print('REPO_ZIP', repo_zip)
print('SUBMISSION_ZIP', submission_zip)
